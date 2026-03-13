# main.py
from fastapi.responses import HTMLResponse
from fastapi import FastAPI, HTTPException, Depends, status, Request, Response, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware import Middleware
from fastapi.middleware.cors import CORSMiddleware
from prometheus_fastapi_instrumentator import Instrumentator
from prometheus_client import Counter, Histogram, Gauge
from starlette.middleware.sessions import SessionMiddleware
import time
import asyncio
import logging
from pydantic import BaseModel
from typing import Optional
import requests
import os
import uuid
import json # <-- AÑADIDO: Para manejar el historial en formato JSON
from authlib.integrations.starlette_client import OAuth
from starlette.config import Config
from jose import JWTError, jwt
from urllib.parse import urlencode, quote_plus 

import httpx
import magic # libreria para determinar el tipo de formato en un archivo a ingestar


# --- IMPORTACIONES DE LANGCHAIN ---
from langchain_ollama import OllamaLLM, OllamaEmbeddings, ChatOllama
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader

# --- AGENT ORCHESTRATOR ---
from agents.state import AgentState
from agents.orchestrator import get_orchestrator
from agents.security import validate_prompt, sanitize_output

# --- MEMORIA Y PERSISTENCIA ---
import redis
import psycopg2
from psycopg2 import pool
from minio import Minio

# ... IMPORTACIONDES DE TENSOFLOW2
import base64
from PIL import Image
import numpy as np
import io

# --- CONFIGURACIÓN DEL SERVICIO DE PRODUCTIVIDAD ---
PRODUCTIVITY_SERVICE_URL  = "http://productivity-service:8001"
K8S_AGENT_URL             = "http://k8s-agent-service:8002"
WHATSAPP_BRIDGE_URL       = os.getenv("WHATSAPP_BRIDGE_URL", "http://whatsapp-bridge-service:3000")
WHATSAPP_PERSONAL_URL     = os.getenv("WHATSAPP_PERSONAL_URL", "http://whatsapp-personal-service:3001")
ADMIN_PHONE               = os.getenv("ADMIN_PHONE", "5219993437008")
INTERNAL_API_SECRET       = os.environ.get("INTERNAL_API_SECRET")

if not INTERNAL_API_SECRET:
    raise ValueError("Falta la variable de entorno 'INTERNAL_API_SECRET' montada en el Pod.")

# URL del nuevo servicio ejecutor dentro del clúster (ELIMINADO)

# --- CONFIGURACIÓN DE OAUTH Y JWT ---
config = Config(environ=os.environ)
oauth = OAuth(config)

oauth.register(
    name='google',
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={
        'scope': 'openid email profile'
    }
)

# OAuth para Google Calendar (offline access + Calendar + Gmail scopes)
oauth.register(
    name='google_calendar',
    client_id=os.environ.get('GOOGLE_CLIENT_ID'),
    client_secret=os.environ.get('GOOGLE_CLIENT_SECRET'),
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={
        'scope': 'openid email https://www.googleapis.com/auth/calendar https://www.googleapis.com/auth/gmail.readonly',
        'access_type': 'offline',
        'prompt': 'consent',
    }
)

#--- LISTA BLANCA DE USUARIOS (desde ConfigMap) ---
allowed_emails_csv = os.environ.get("ALLOWED_EMAILS_CSV", "")
ALLOWED_EMAILS = [email.strip() for email in allowed_emails_csv.split(',') if email.strip()]

allowed_numbers_csv = os.environ.get("ALLOWED_NUMBERS_CSV", "")
ALLOWED_NUMBERS = [num.strip() for num in allowed_numbers_csv.split(',') if num.strip()]

# Combinar ambas para la whitelist general del chat
FULL_WHITELIST = ALLOWED_EMAILS + ALLOWED_NUMBERS

k8s_allowed_csv = os.environ.get("K8S_ALLOWED_USERS_CSV", "")
K8S_ALLOWED_USERS = [u.strip() for u in k8s_allowed_csv.split(',') if u.strip()]

if not K8S_ALLOWED_USERS:
    print("Warning: No se han configurado usuarios autorizados para K8s.")

if not K8S_ALLOWED_USERS:
    print("Warning: No se han configurado usuarios autorizados para K8s. Usando FULL_WHITELIST como respaldo temporal.")
    K8S_ALLOWED_USERS = FULL_WHITELIST


# --- FUNCIONES DE AUTENTICACIÓN JWT ---
SECRET_KEY = os.environ.get("JWT_SECRET_KEY")
ALGORITHM = "HS256"

security = HTTPBearer()

def create_jwt_token(email: str):
    to_encode = {"sub": email}
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None or email not in FULL_WHITELIST:
            raise HTTPException(status_code=403, detail="Usuario no autorizado")
        return email
    except JWTError:
        raise HTTPException(status_code=401, detail="Token inválido o expirado")

# --- CONFIGURACIÓN E INICIALIZACIÓN DE IA ---
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://ollama-service:11434")
LLM_MODEL = os.getenv("LLM_MODEL", "qwen2.5:14b")       # Modelo principal
VISION_MODEL = os.getenv("LLM_VISION_MODEL", "qwen2.5-vl:7b") # Modelo Vision
EMBED_MODEL = "nomic-embed-text" 

# Modelos definitivos
llm = OllamaLLM(model=LLM_MODEL, base_url=OLLAMA_BASE_URL)
vision_llm = ChatOllama(model=VISION_MODEL, base_url=OLLAMA_BASE_URL)

# --- CONFIGURACIÓN DE CAPA DE DATOS (PHASE 2) ---
REDIS_HOST = os.getenv("REDIS_HOST", "redis-service")
REDIS_PORT = int(os.getenv("REDIS_PORT", 6379))
redis_client = redis.Redis(host=REDIS_HOST, port=REDIS_PORT, decode_responses=True)

# P4: Rate limiting — max requests per user per window
RATE_LIMIT_MAX = int(os.getenv("RATE_LIMIT_MAX", 15))       # requests
RATE_LIMIT_WINDOW = int(os.getenv("RATE_LIMIT_WINDOW", 60)) # seconds

def _check_rate_limit(user_id: str) -> tuple[bool, int]:
    """
    Returns (allowed, requests_remaining).
    Uses Redis incr+expire for a sliding fixed-window counter.
    """
    key = f"rate_limit:{user_id}"
    try:
        current = redis_client.incr(key)
        if current == 1:
            redis_client.expire(key, RATE_LIMIT_WINDOW)
        remaining = max(0, RATE_LIMIT_MAX - current)
        return current <= RATE_LIMIT_MAX, remaining
    except Exception as exc:
        logging.warning(f"[RATE_LIMIT] Redis error, allowing request: {exc}")
        return True, RATE_LIMIT_MAX

QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant-service:6333")

# P5-3: Module-level singletons for Qdrant and embeddings (avoid per-request reconnections)
_qdrant_client = QdrantClient(url=QDRANT_URL)
_embeddings = OllamaEmbeddings(model=EMBED_MODEL, base_url=OLLAMA_BASE_URL)

# P5-4: Persistent httpx client with connection pooling for internal service calls
_http_client = httpx.Client(timeout=120.0)

# Module-level LLM for document generation (longer context, higher quality)
from langchain_ollama import OllamaLLM as _OllamaLLM
_llm_doc = _OllamaLLM(model=LLM_MODEL, base_url=OLLAMA_BASE_URL, num_predict=2000)

POSTGRES_HOST = os.getenv("POSTGRES_HOST", "postgres-service")
POSTGRES_DB = os.getenv("POSTGRES_DB", "amael_db")
POSTGRES_USER = os.getenv("POSTGRES_USER", "amael_user")
POSTGRES_PASS = os.getenv("POSTGRES_PASSWORD", "amael_password_2026")

# Pool de conexiones para Postgres
_postgres_pool = None
_last_pool_attempt = 0
POOL_RETRY_INTERVAL = 30 # seconds

def get_postgres_pool():
    global _postgres_pool, _last_pool_attempt
    now = time.time()
    
    if _postgres_pool is not None:
        return _postgres_pool
    
    if now - _last_pool_attempt < POOL_RETRY_INTERVAL:
        return None
        
    _last_pool_attempt = now
    try:
        _postgres_pool = psycopg2.pool.SimpleConnectionPool(1, 10,
            user=POSTGRES_USER,
            password=POSTGRES_PASS,
            host=POSTGRES_HOST,
            database=POSTGRES_DB
        )
        print("PostgreSQL connection pool created successfully")
        return _postgres_pool
    except Exception as e:
        print(f"Error creating PostgreSQL pool: {e}")
        return None

MINIO_URL = os.getenv("MINIO_URL", "minio-service:9000")
MINIO_ACCESS_KEY = os.getenv("MINIO_ACCESS_KEY", "amael_admin")
MINIO_SECRET_KEY = os.getenv("MINIO_SECRET_KEY", "amael_minio_secret_key")

minio_client = Minio(
    MINIO_URL,
    access_key=MINIO_ACCESS_KEY,
    secret_key=MINIO_SECRET_KEY,
    secure=False
)
CHAT_HISTORIES_DIR = os.getenv("CHAT_HISTORIES_DIR", "/chat_histories")

# --- CUSTOM PROMETHEUS METRICS ---
LLM_TOKENS_TOTAL = Counter('amael_llm_tokens_total', 'Total tokens used by LLM', ['model', 'type'])
LLM_LATENCY_SECONDS = Histogram('amael_llm_latency_seconds', 'Latency of LLM requests in seconds', ['model'])
AGENT_STEPS_TOTAL = Counter('amael_agent_steps_total', 'Total steps taken by the agent')
AGENT_FAILURES_TOTAL = Counter('amael_agent_failures_total', 'Total failures in agent execution')
AGENT_TOOLS_USAGE_TOTAL = Counter('amael_agent_tools_usage_total', 'Total usage of agent tools', ['tool'])
RAG_HITS_TOTAL = Counter('amael_rag_hits_total', 'Total number of RAG hits')
RAG_MISS_TOTAL = Counter('amael_rag_miss_total', 'Total number of RAG misses')

# --- AGENT METRICS ---
PLANNER_STEPS_TOTAL = Counter('amael_planner_steps_total', 'Total steps generated by the planner')
AGENT_EXECUTION_LATENCY = Histogram('amael_agent_execution_latency', 'Latency of agent execution in seconds', buckets=(0.5, 1, 2, 5, 10, 20, 40, 60, 120))
TOOL_CALLS_TOTAL = Counter('amael_tool_calls_total', 'Total number of tool calls', ['tool'])
REASONING_ITERATIONS = Counter('amael_reasoning_iterations', 'Total iterations of reasoning')
# P6: Security observability
SECURITY_RATE_LIMITED_TOTAL = Counter('amael_security_rate_limited_total', 'Requests blocked by rate limiter')
SECURITY_INPUT_BLOCKED_TOTAL = Counter('amael_security_input_blocked_total', 'Prompts rejected by input validation', ['reason'])

# --- DEFINICIÓN DEL PROMPT (MOVIDO AQUÍ - ANTES DE LAS FUNCIONES QUE LO USAN) ---
system_prompt_template = """
### PERSONAJE
Eres un asistente de IA avanzado con nombre de Amael creado por Ricardo Guzman, versátil y servicial. Tu objetivo es proporcionar respuestas precisas, claras y útiles. Adapta tu estilo de respuesta a la naturaleza de la pregunta del usuario, adaptate a su lenguaje.

### REGLAS DE INTERACCIÓN
    Usa el Historial y el Contexto: Analiza primero el HISTORIAL DE LA CONVERSACIÓN y luego el CONTEXTO DE DOCUMENTOS. 
    Sé Natural: Para preguntas simples, responde de forma natural. 
    Sé Transparente: Si el CONTEXTO no contiene la información, pero tu conocimiento general te permite responder, indícalo explícitamente. 
    No Inventes Datos Específicos: Nunca inventes métricas, nombres de archivos o detalles que no estén en el CONTEXTO o el HISTORIAL. 

### INSTRUCCIONES
1.  **Analiza la Petición del Usuario:** Entiende la pregunta o solicitud del usuario.
2.  **Usa el Contexto Disponible:**
    - Primero, revisa el `HISTORIAL DE LA CONVERSACIÓN` para entender el contexto.
    - Luego, utiliza el `CONTEXTO DE DOCUMENTOS` para fundamentar tu respuesta con datos específicos.
3.  **Decide Cómo Responder:**
    - **¿Es una pregunta general o conversacional?** (ej: "hola", "¿cómo estás?"). Responde de forma natural y directa.
    - **¿Es una pregunta compleja?** Estructura tu respuesta con Markdown (títulos, listas, negrita) para que sea clara y fácil de leer.
4.  **Sé Preciso y Transparente:** Basa tus respuestas en la información proporcionada. Si el contexto es insuficiente, admítelo. Si usas tu conocimiento general, indícalo.
5.  **No Inventes Datos Específicos:** Nunca inventes métricas, nombres de archivo o detalles que no estén en el `CONTEXTO` o el `HISTORIAL`.

### HISTORIAL DE LA CONVERSACIÓN
---
{conversation_history}
---

### CONTEXTO DE DOCUMENTOS
---
<<<CONTEXT>>>

**Pregunta del Usuario:**
{request_prompt}

**Respuesta del Asistente:**
"""

# --- FUNCIONES AUXILIARES MULTIUSUARIO ---

def sanitize_email(email: str) -> str:
    """Crea un nombre de directorio seguro a partir de un email."""
    return email.replace("@", "_at_").replace(".", "_dot_")

def get_user_vectorstore(user_email: str):
    """Carga o crea la base de datos vectorial para un usuario específico en Qdrant."""
    collection_name = sanitize_email(user_email)

    try:
        vectorstore = QdrantVectorStore.from_existing_collection(
            embedding=_embeddings,
            collection_name=collection_name,
            url=QDRANT_URL,
        )
        return vectorstore
    except Exception as e:
        print(f"[Qdrant] Error o diferencia de dimensiones en {collection_name}: {e}. Recreando...")

        # Si la colección existe pero dio error (probablemente dimensiones), la borramos
        if _qdrant_client.collection_exists(collection_name):
            print(f"[Qdrant] Borrando colección existente {collection_name} para corregir configuración...")
            _qdrant_client.delete_collection(collection_name)

        _qdrant_client.create_collection(
            collection_name=collection_name,
            vectors_config={"size": 768, "distance": "Cosine"} # nomic-embed-text es de 768
        )

        return QdrantVectorStore(
            client=_qdrant_client,
            collection_name=collection_name,
            embedding=_embeddings
        )

def init_db():
    """Inicializa las tablas necesarias en PostgreSQL."""
    pool = get_postgres_pool()
    if not pool: return
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS chat_history (
                    id SERIAL PRIMARY KEY,
                    user_id TEXT NOT NULL,
                    role TEXT NOT NULL,
                    content TEXT NOT NULL,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );
                CREATE INDEX IF NOT EXISTS idx_user_id ON chat_history(user_id);
                
                CREATE TABLE IF NOT EXISTS user_documents (
                    id SERIAL PRIMARY KEY,
                    user_id TEXT NOT NULL,
                    doc_type TEXT,
                    summary TEXT,
                    content JSONB,
                    raw_analysis TEXT,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );
                CREATE INDEX IF NOT EXISTS idx_user_docs_user_id ON user_documents(user_id);

                CREATE TABLE IF NOT EXISTS conversations (
                    id SERIAL PRIMARY KEY,
                    user_id TEXT NOT NULL,
                    title TEXT DEFAULT 'Nueva conversación',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    last_active_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );
                CREATE INDEX IF NOT EXISTS idx_conv_user_id ON conversations(user_id);

                CREATE TABLE IF NOT EXISTS message_feedback (
                    id SERIAL PRIMARY KEY,
                    user_id TEXT NOT NULL,
                    conversation_id INTEGER,
                    message_index INTEGER,
                    sentiment TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS user_profile (
                    user_id TEXT PRIMARY KEY,
                    display_name TEXT,
                    timezone TEXT DEFAULT 'America/Mexico_City',
                    preferences JSONB DEFAULT '{}',
                    context_data JSONB DEFAULT '{}',
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );

                CREATE TABLE IF NOT EXISTS user_facts (
                    id SERIAL PRIMARY KEY,
                    user_id TEXT NOT NULL,
                    fact TEXT NOT NULL,
                    category TEXT DEFAULT 'general',
                    source_conv_id INTEGER,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );
                CREATE INDEX IF NOT EXISTS idx_user_facts_uid ON user_facts(user_id);

                CREATE TABLE IF NOT EXISTS user_goals (
                    id SERIAL PRIMARY KEY,
                    user_id TEXT NOT NULL,
                    title TEXT NOT NULL,
                    category TEXT DEFAULT 'personal',
                    status TEXT DEFAULT 'active',
                    progress INTEGER DEFAULT 0,
                    deadline DATE,
                    milestones JSONB DEFAULT '[]',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );
                CREATE INDEX IF NOT EXISTS idx_user_goals_uid ON user_goals(user_id);
            """)
            # Add conversation_id column to chat_history if it doesn't exist yet
            cur.execute("""
                ALTER TABLE chat_history ADD COLUMN IF NOT EXISTS conversation_id INTEGER;
            """)
            # Add role + status to user_profile
            cur.execute("""
                ALTER TABLE user_profile ADD COLUMN IF NOT EXISTS role VARCHAR DEFAULT 'user';
                ALTER TABLE user_profile ADD COLUMN IF NOT EXISTS status VARCHAR DEFAULT 'active';
                ALTER TABLE user_profile ADD COLUMN IF NOT EXISTS timezone TEXT DEFAULT 'America/Mexico_City';
            """)
            # User identities: maps phone/email → canonical user_id
            cur.execute("""
                CREATE TABLE IF NOT EXISTS user_identities (
                    id SERIAL PRIMARY KEY,
                    canonical_user_id TEXT NOT NULL,
                    identity_type TEXT NOT NULL,
                    identity_value TEXT UNIQUE NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );
                CREATE INDEX IF NOT EXISTS idx_identity_value ON user_identities(identity_value);
            """)
            # WhatsApp Personal: settings por usuario (horarios, reglas)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS whatsapp_personal_settings (
                    user_id       TEXT PRIMARY KEY,
                    auto_reply    BOOLEAN     DEFAULT TRUE,
                    quiet_start   INTEGER     DEFAULT 22,
                    quiet_end     INTEGER     DEFAULT 8,
                    active_days   INTEGER[]   DEFAULT '{1,2,3,4,5}',
                    reply_scope   TEXT        DEFAULT 'all',
                    ai_assist     BOOLEAN     DEFAULT TRUE,
                    offline_msg       TEXT        DEFAULT NULL,
                    updated_at        TIMESTAMP   DEFAULT NOW()
                );
                ALTER TABLE whatsapp_personal_settings
                    ADD COLUMN IF NOT EXISTS quiet_enabled    BOOLEAN DEFAULT TRUE;
                ALTER TABLE whatsapp_personal_settings
                    ADD COLUMN IF NOT EXISTS allowed_contacts JSONB   DEFAULT '[]';
            """)
            # Platform-wide settings (key/value)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS platform_settings (
                    key TEXT PRIMARY KEY,
                    value TEXT NOT NULL,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );
                INSERT INTO platform_settings (key, value) VALUES ('allow_access_requests', 'false')
                ON CONFLICT (key) DO NOTHING;
            """)
            # Seed admin user + identities
            _admin_email = os.environ.get('ADMIN_EMAIL', 'ricardogs26@gmail.com')
            _admin_phone = os.environ.get('ADMIN_PHONE', '5219993437008')
            cur.execute("""
                INSERT INTO user_profile (user_id, role, status)
                VALUES (%s, 'admin', 'active')
                ON CONFLICT (user_id) DO UPDATE SET role = 'admin', status = 'active'
                WHERE user_profile.role != 'admin';
            """, (_admin_email,))
            cur.execute("""
                INSERT INTO user_identities (canonical_user_id, identity_type, identity_value)
                VALUES (%s, 'email', %s) ON CONFLICT (identity_value) DO NOTHING;
            """, (_admin_email, _admin_email))
            if _admin_phone:
                cur.execute("""
                    INSERT INTO user_identities (canonical_user_id, identity_type, identity_value)
                    VALUES (%s, 'whatsapp', %s) ON CONFLICT (identity_value) DO NOTHING;
                """, (_admin_email, _admin_phone))
            conn.commit()
    finally:
        pool.putconn(conn)


# ─── Identity helpers ──────────────────────────────────────────────────────────

def resolve_user_id(identifier: str) -> str:
    """Mapea cualquier identidad (teléfono/email) al user_id canónico (email)."""
    pool = get_postgres_pool()
    if not pool or not identifier:
        return identifier
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT canonical_user_id FROM user_identities WHERE identity_value = %s", (identifier,))
            row = cur.fetchone()
            return row[0] if row else identifier
    finally:
        pool.putconn(conn)


def get_user_role(user_id: str) -> str:
    """Devuelve el rol del usuario: 'admin', 'user' o 'readonly'."""
    pool = get_postgres_pool()
    if not pool:
        return 'admin' if user_id == os.environ.get('ADMIN_EMAIL', 'ricardogs26@gmail.com') else 'user'
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT role FROM user_profile WHERE user_id = %s", (user_id,))
            row = cur.fetchone()
            return row[0] if row else 'user'
    finally:
        pool.putconn(conn)


def is_user_registered(identifier: str):
    """Retorna (allowed: bool, canonical_user_id: str)."""
    canonical = resolve_user_id(identifier)
    if identifier in FULL_WHITELIST or canonical in FULL_WHITELIST:
        return True, canonical
    pool = get_postgres_pool()
    if not pool:
        return False, canonical
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT user_id FROM user_profile WHERE user_id = %s AND status = 'active'",
                (canonical,)
            )
            return (cur.fetchone() is not None, canonical)
    finally:
        pool.putconn(conn)


def get_platform_setting(key: str, default: str = '') -> str:
    pool = get_postgres_pool()
    if not pool:
        return default
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT value FROM platform_settings WHERE key = %s", (key,))
            row = cur.fetchone()
            return row[0] if row else default
    finally:
        pool.putconn(conn)


# ───────────────────────────────────────────────────────────────────────────────

def save_chat_message(user_id: str, role: str, content: str, conversation_id: Optional[int] = None):
    """Guarda un mensaje en Redis (caché) y Postgres (persistente)."""
    if content is None:
        content = ""

    redis_key = f"chat_history:{user_id}:{conversation_id}" if conversation_id else f"chat_history:{user_id}"
    message_json = json.dumps({"role": role, "content": content})
    redis_client.lpush(redis_key, message_json)
    redis_client.ltrim(redis_key, 0, 19)

    pool = get_postgres_pool()
    if pool:
        conn = pool.getconn()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "INSERT INTO chat_history (user_id, role, content, conversation_id) VALUES (%s, %s, %s, %s)",
                    (user_id, role, content, conversation_id)
                )
                conn.commit()
        finally:
            pool.putconn(conn)

    if conversation_id:
        _touch_conversation(conversation_id)


def _touch_conversation(conversation_id: int):
    pool = get_postgres_pool()
    if not pool: return
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("UPDATE conversations SET last_active_at = NOW() WHERE id = %s", (conversation_id,))
            conn.commit()
    finally:
        pool.putconn(conn)


def create_conversation(user_id: str, title: str = "Nueva conversación") -> Optional[int]:
    pool = get_postgres_pool()
    if not pool: return None
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO conversations (user_id, title) VALUES (%s, %s) RETURNING id",
                (user_id, title)
            )
            conv_id = cur.fetchone()[0]
            conn.commit()
            return conv_id
    finally:
        pool.putconn(conn)


def get_user_conversations(user_id: str, limit: int = 30) -> list:
    pool = get_postgres_pool()
    if not pool: return []
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, title, created_at, last_active_at FROM conversations WHERE user_id = %s ORDER BY last_active_at DESC LIMIT %s",
                (user_id, limit)
            )
            rows = cur.fetchall()
            return [{"id": r[0], "title": r[1], "created_at": r[2].isoformat(), "last_active_at": r[3].isoformat()} for r in rows]
    finally:
        pool.putconn(conn)


def get_conversation_messages(conversation_id: int, user_id: str) -> list:
    pool = get_postgres_pool()
    if not pool: return []
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT role, content, timestamp FROM chat_history WHERE conversation_id = %s AND user_id = %s ORDER BY timestamp ASC",
                (conversation_id, user_id)
            )
            rows = cur.fetchall()
            return [{"role": r[0], "content": r[1], "ts": r[2].strftime("%H:%M")} for r in rows]
    finally:
        pool.putconn(conn)


def get_chat_history(user_id: str, limit: int = 20, conversation_id: Optional[int] = None) -> list:
    """Intenta obtener el historial de Redis, si falla va a Postgres."""
    redis_key = f"chat_history:{user_id}:{conversation_id}" if conversation_id else f"chat_history:{user_id}"
    try:
        history = redis_client.lrange(redis_key, 0, limit - 1)
        if history:
            return [json.loads(m) for m in reversed(history)]
    except Exception as e:
        print(f"Redis error: {e}")

    pool = get_postgres_pool()
    if pool:
        conn = pool.getconn()
        try:
            with conn.cursor() as cur:
                if conversation_id:
                    cur.execute(
                        "SELECT role, content FROM chat_history WHERE user_id = %s AND conversation_id = %s ORDER BY timestamp DESC LIMIT %s",
                        (user_id, conversation_id, limit)
                    )
                else:
                    cur.execute(
                        "SELECT role, content FROM chat_history WHERE user_id = %s AND conversation_id IS NULL ORDER BY timestamp DESC LIMIT %s",
                        (user_id, limit)
                    )
                rows = cur.fetchall()
                return [{"role": r, "content": c} for r, c in reversed(rows)]
        finally:
            pool.putconn(conn)
    
    return []

def save_user_document(user_id: str, doc_type: str, summary: str, content: dict, raw_analysis: str):
    """Guarda un documento extraído (como un ticket) en Postgres."""
    pool = get_postgres_pool()
    if pool:
        conn = pool.getconn()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "INSERT INTO user_documents (user_id, doc_type, summary, content, raw_analysis) VALUES (%s, %s, %s, %s, %s)",
                    (user_id, doc_type, summary, json.dumps(content), raw_analysis)
                )
                conn.commit()
        finally:
            pool.putconn(conn)

def get_user_documents(user_id: str, query: str = None, limit: int = 5):
    """Recupera documentos guardados de un usuario."""
    pool = get_postgres_pool()
    if pool:
        conn = pool.getconn()
        try:
            with conn.cursor() as cur:
                if query:
                    cur.execute(
                        "SELECT doc_type, summary, timestamp FROM user_documents WHERE user_id = %s AND (summary ILIKE %s OR doc_type ILIKE %s) ORDER BY timestamp DESC LIMIT %s",
                        (user_id, f"%{query}%", f"%{query}%", limit)
                    )
                else:
                    cur.execute(
                        "SELECT doc_type, summary, timestamp FROM user_documents WHERE user_id = %s ORDER BY timestamp DESC LIMIT %s",
                        (user_id, limit)
                    )
                return cur.fetchall()
        finally:
            pool.putconn(conn)
    return []

def get_history_path_for_id(identifier: str) -> str:
    """Obtiene la ruta al archivo de historial para un identificador único (email o teléfono)."""
    # Sanitiza el identificador para que sea un nombre de archivo válido
    sanitized_id = identifier.replace("@", "_at_").replace(".", "_dot_").replace("-", "_dash_")
    return os.path.join(CHAT_HISTORIES_DIR, f"{sanitized_id}.json")

# --- MODELOS DE DATOS ---
class ChatRequest(BaseModel):
    prompt: str
    history: list[dict] = []
    user_id: str = None
    image: str = None
    conversation_id: Optional[int] = None

class ChatResponse(BaseModel):
    response: str

class ConversationCreate(BaseModel):
    title: str = "Nueva conversación"

class FeedbackRequest(BaseModel):
    conversation_id: Optional[int] = None
    message_index: int
    sentiment: str  # 'positive' or 'negative'

class SREKnowledgeRequest(BaseModel):
    key: str
    content: str

class SREQueryRequest(BaseModel):
    query: str

# --- CREACIÓN DE LA APLICACIÓN FASTAPI CON MIDDLEWARE ---
app = FastAPI(
    title="API Agente Personal",
    version="2.1",
    middleware=[
        Middleware(SessionMiddleware, secret_key=os.environ.get("SESSION_SECRET_KEY")),
        Middleware(
            CORSMiddleware,
            allow_origins=["*"], # En producción, restringe esto a tu dominio
            allow_credentials=True,
            allow_methods=["*"],
            allow_headers=["*"],
        ),
    ]
)

# Prometheus HTTP metrics
Instrumentator().instrument(app).expose(app)

@app.on_event("startup")
async def startup_event():
    init_db()
    # OpenTelemetry: auto-instrument FastAPI + HTTPX outbound calls
    from agents.tracing import instrument_app
    instrument_app(app)

# --- ENDPOINTS DE AUTENTICACIÓN (Sin cambios) ---
@app.get('/api/auth/login')
async def login(request: Request):
    redirect_uri = "https://amael-ia.richardx.dev/api/auth/callback"
    return await oauth.google.authorize_redirect(request, redirect_uri)

@app.get('/api/auth/callback')
async def auth_callback(request: Request):
    try:
        token = await oauth.google.authorize_access_token(request)
        user_info = token.get('userinfo')
        if user_info and user_info['email'] in ALLOWED_EMAILS:
            jwt_token = create_jwt_token(user_info['email'])
            frontend_url = "https://amael-ia.richardx.dev"
            
            # --- CAMBIO IMPORTANTE AQUÍ ---
            # Extraemos la información del perfil
            user_name = user_info.get('name', 'Usuario')
            user_picture = user_info.get('picture')

            # Creamos un diccionario con los parámetros
            params = {
                "token": jwt_token,
                "name": user_name,
                "picture": user_picture
            }
            
            # Codificamos los parámetros para la URL
            redirect_url = f"{frontend_url}?{urlencode(params)}"
            
            # Redirigimos con el token y la información del perfil
            return Response(status_code=302, headers={"location": redirect_url})
        else:
            # Manejo de error también con redirección
            frontend_url = "https://amael-ia.richardx.dev"
            return Response(status_code=302, headers={"location": f"{frontend_url}?error=unauthorized"})

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en la autenticación con Google: {e}")

# --- GOOGLE CALENDAR OAUTH ---

@app.get('/api/auth/calendar')
async def calendar_auth(request: Request, token: Optional[str] = None):
    """Inicia el flujo OAuth para autorizar acceso a Google Calendar y Gmail.
    Acepta el JWT como query param ?token= (navegación directa del browser)."""
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if not email or email not in FULL_WHITELIST:
            raise HTTPException(status_code=403, detail="Usuario no autorizado")
    except JWTError:
        raise HTTPException(status_code=401, detail="Token inválido")
    redirect_uri = "https://amael-ia.richardx.dev/api/auth/calendar/callback"
    request.session['calendar_user'] = email
    return await oauth.google_calendar.authorize_redirect(
        request, redirect_uri,
        access_type='offline',
        prompt='consent',
    )


@app.get('/api/auth/calendar/callback')
async def calendar_callback(request: Request):
    """Recibe el código OAuth, obtiene los tokens y los guarda en Vault via productivity-service."""
    try:
        token = await oauth.google_calendar.authorize_access_token(request)
        user_email = request.session.get('calendar_user')
        if not user_email:
            user_info = token.get('userinfo', {})
            user_email = user_info.get('email', '')

        if not user_email:
            raise HTTPException(status_code=400, detail="No se pudo identificar el usuario.")

        refresh_token = token.get('refresh_token')
        if not refresh_token:
            logging.error(f"[CALENDAR] Google no devolvió refresh_token para {user_email}. Token keys: {list(token.keys())}")
            return Response(status_code=302, headers={"location": "https://amael-ia.richardx.dev?calendar_error=1"})

        # Guardar tokens en Vault via productivity-service
        payload = {
            "user_email": user_email,
            "token": token.get('access_token'),
            "refresh_token": refresh_token,
            "token_uri": "https://oauth2.googleapis.com/token",
            "client_id": os.environ.get('GOOGLE_CLIENT_ID'),
            "client_secret": os.environ.get('GOOGLE_CLIENT_SECRET'),
            "scopes": ["https://www.googleapis.com/auth/calendar", "https://www.googleapis.com/auth/gmail.readonly"],
        }
        r = _http_client.post(
            f"{PRODUCTIVITY_SERVICE_URL}/credentials",
            json=payload,
            headers={"Authorization": f"Bearer {INTERNAL_API_SECRET}"},
            timeout=30.0,
        )
        if r.status_code != 200:
            logging.error(f"[CALENDAR] Error guardando credenciales: {r.text}")
            return Response(status_code=302, headers={"location": "https://amael-ia.richardx.dev?calendar_error=1"})

        logging.info(f"[CALENDAR] Credenciales guardadas para {user_email}")
        return Response(status_code=302, headers={"location": "https://amael-ia.richardx.dev?calendar_connected=1"})

    except Exception as e:
        logging.error(f"[CALENDAR] Error en callback: {e}")
        return Response(status_code=302, headers={"location": "https://amael-ia.richardx.dev?calendar_error=1"})


@app.get('/api/auth/calendar/status')
async def calendar_status(user: str = Depends(get_current_user)):
    """Verifica si el usuario tiene Google Calendar conectado."""
    r = _http_client.get(
        f"{PRODUCTIVITY_SERVICE_URL}/credentials/status",
        params={"user_email": user},
        headers={"Authorization": f"Bearer {INTERNAL_API_SECRET}"},
        timeout=10.0,
    )
    return r.json() if r.status_code == 200 else {"connected": False}


# --- ENDPOINTS DE LA APLICACIÓN (ahora protegidos y multiusuario) ---

# Currencies/tickers detected → use exchange rate API instead of DDG
_CURRENCY_KEYWORDS = {"dolar", "dollar", "usd", "euro", "eur", "tipo de cambio",
                       "precio dolar", "cotizacion", "cotización", "divisas", "forex"}

def _web_search(query: str) -> str:
    """Web search: exchange rate API for currency queries, DuckDuckGo for the rest."""
    q_lower = query.lower()

    # ── Fast path: currency rates via free API ────────────────────────────────
    if any(kw in q_lower for kw in _CURRENCY_KEYWORDS):
        try:
            r = _http_client.get("https://open.er-api.com/v6/latest/USD", timeout=10.0)
            if r.status_code == 200:
                data = r.json()
                rates = data.get("rates", {})
                updated = data.get("time_last_update_utc", "")
                mxn = rates.get("MXN")
                eur = rates.get("EUR")
                cad = rates.get("CAD")
                gbp = rates.get("GBP")
                lines = ["**Tipo de cambio actual (base USD):**"]
                if mxn: lines.append(f"• 1 USD = **{mxn:.4f} MXN** (Peso Mexicano)")
                if eur: lines.append(f"• 1 USD = **{eur:.4f} EUR** (Euro)")
                if cad: lines.append(f"• 1 USD = **{cad:.4f} CAD** (Dólar Canadiense)")
                if gbp: lines.append(f"• 1 USD = **{gbp:.4f} GBP** (Libra Esterlina)")
                lines.append(f"\nActualizado: {updated}\nFuente: open.er-api.com")
                return "\n".join(lines)
        except Exception as e:
            logger.warning(f"[WEB_SEARCH] Exchange rate API falló: {e}")

    # ── General search via DuckDuckGo ─────────────────────────────────────────
    try:
        from duckduckgo_search import DDGS
        with DDGS() as ddgs:
            results = list(ddgs.text(query, region="mx-es", safesearch="off",
                                     timelimit="m", max_results=5))
        if not results:
            return "No se encontraron resultados para la búsqueda."
        lines = []
        for r in results:
            title = r.get("title", "")
            body = r.get("body", "")
            href = r.get("href", "")
            lines.append(f"**{title}**\n{body}\nFuente: {href}")
        return "\n\n---\n\n".join(lines)
    except Exception as e:
        return f"Error en búsqueda web: {e}"


def _extract_docx_text(content: bytes) -> str:
    """Extrae texto plano de un archivo DOCX."""
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


@app.post("/api/ingest")
async def ingest_data(file: UploadFile = File(...), user: str = Depends(get_current_user)):
    """Sube y procesa documentos (PDF, TXT, DOCX) — indexa en Qdrant y guarda metadata en DB."""
    if get_user_role(user) == 'readonly':
        raise HTTPException(status_code=403, detail="Tu rol de solo lectura no permite subir documentos.")
    temp_file_path = f"/tmp/{uuid.uuid4()}-{file.filename}"
    try:
        content = await file.read()

        # Detectar MIME type
        try:
            mime = magic.from_buffer(content, mime=True)
        except Exception:
            raise HTTPException(status_code=400, detail="No se pudo determinar el tipo de archivo.")

        # Soporte DOCX por extensión (magic lo detecta como zip)
        fname_lower = (file.filename or "").lower()
        if fname_lower.endswith(".docx"):
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # Extraer texto y chunks según tipo
        full_text = ""
        from langchain_core.documents import Document as LCDocument

        if mime == "application/pdf":
            with open(temp_file_path, "wb") as buf:
                buf.write(content)
            loader = PyPDFLoader(temp_file_path)
            documents = loader.load()
            full_text = "\n".join(d.page_content for d in documents)
        elif mime == "text/plain":
            full_text = content.decode("utf-8", errors="replace")
            documents = [LCDocument(page_content=full_text)]
        elif "wordprocessingml" in mime or fname_lower.endswith(".docx"):
            full_text = _extract_docx_text(content)
            documents = [LCDocument(page_content=full_text)]
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Tipo de archivo no soportado: '{mime}'. Usa PDF, TXT o DOCX."
            )

        # Chunking + indexar en Qdrant
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        texts = splitter.split_documents(documents)
        user_vectorstore = get_user_vectorstore(user)
        user_vectorstore.add_documents(texts)

        # Generar resumen con LLM (primeros 3000 chars)
        summary = ""
        try:
            llm_sum = OllamaLLM(model=os.environ.get("MODEL_NAME", "qwen2.5:14b"),
                                base_url=os.environ.get("OLLAMA_BASE_URL", "http://ollama-service:11434"))
            summary = llm_sum.invoke(
                f"Resume el siguiente documento en 2-3 oraciones en español:\n\n{full_text[:3000]}"
            ).strip()
        except Exception as e:
            summary = f"Documento procesado ({len(texts)} fragmentos)."
            logger.warning(f"[INGEST] Error generando summary: {e}")

        # Guardar metadata en user_documents
        doc_id = None
        pool = get_postgres_pool()
        if pool:
            conn = pool.getconn()
            try:
                with conn.cursor() as cur:
                    cur.execute(
                        "INSERT INTO user_documents (user_id, doc_type, summary, raw_analysis) "
                        "VALUES (%s, %s, %s, %s) RETURNING id",
                        (user, file.filename, summary, full_text[:10000])
                    )
                    doc_id = cur.fetchone()[0]
                    conn.commit()
            finally:
                pool.putconn(conn)

        # Backup en MinIO (best-effort)
        try:
            bucket_name = sanitize_email(user).replace("_", "-")
            if not minio_client.bucket_exists(bucket_name):
                minio_client.make_bucket(bucket_name)
            with open(temp_file_path, "wb") as buf:
                buf.write(content)
            with open(temp_file_path, "rb") as f:
                minio_client.put_object(bucket_name, file.filename, f,
                                        length=len(content), content_type=mime)
        except Exception as e:
            logger.warning(f"[INGEST] MinIO backup falló (no crítico): {e}")

        return {
            "doc_id": doc_id,
            "filename": file.filename,
            "summary": summary,
            "chunks": len(texts),
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error procesando el archivo: {e}")
    finally:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)


@app.get("/api/documents")
async def list_documents(user: str = Depends(get_current_user)):
    """Lista los documentos subidos por el usuario."""
    pool = get_postgres_pool()
    if not pool:
        return {"documents": []}
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, doc_type, summary, timestamp FROM user_documents "
                "WHERE user_id = %s ORDER BY timestamp DESC LIMIT 50",
                (user,)
            )
            rows = cur.fetchall()
        return {"documents": [
            {"id": r[0], "filename": r[1], "summary": r[2],
             "created_at": r[3].isoformat() if r[3] else None}
            for r in rows
        ]}
    finally:
        pool.putconn(conn)


@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: int, user: str = Depends(get_current_user)):
    """Elimina un documento del historial (DB). Los vectores en Qdrant se mantienen."""
    pool = get_postgres_pool()
    if not pool:
        raise HTTPException(status_code=503, detail="DB no disponible")
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "DELETE FROM user_documents WHERE id = %s AND user_id = %s RETURNING id",
                (doc_id, user)
            )
            deleted = cur.fetchone()
            conn.commit()
        if not deleted:
            raise HTTPException(status_code=404, detail="Documento no encontrado")
        return {"deleted": doc_id}
    finally:
        pool.putconn(conn)

@app.post("/api/chat", response_model=ChatResponse, dependencies=[Depends(get_current_user)])
async def chat_endpoint(request: ChatRequest, user: str = Depends(get_current_user)):
    """Endpoint para chatear, ahora con soporte para múltiples usuarios por ID y ejecución de herramientas."""
    
    # --- LÓGICA DEL ENDPOINT ---
     
    # Validar que el prompt no esté vacío o contenga solo espacios en blanco.
    if not request.prompt or request.prompt.strip() == "":
        return ChatResponse(response="¡Hola! como te encuentras, envía tu consulta.")

    # P4-2: Rate limiting por usuario
    effective_user = resolve_user_id(request.user_id) if request.user_id else user
    allowed, remaining = _check_rate_limit(effective_user)
    if not allowed:
        logging.warning(f"[SECURITY] Rate limit excedido para user={effective_user}")
        SECURITY_RATE_LIMITED_TOTAL.inc()
        raise HTTPException(
            status_code=429,
            detail=f"Demasiadas solicitudes. Espera {RATE_LIMIT_WINDOW} segundos antes de volver a intentarlo.",
            headers={"Retry-After": str(RATE_LIMIT_WINDOW)},
        )

    # P4-3: Validación y sanitización del prompt de entrada
    is_valid, prompt_or_error = validate_prompt(request.prompt)
    if not is_valid:
        logging.warning(f"[SECURITY] Prompt rechazado para user={effective_user}: {prompt_or_error}")
        if len(request.prompt) > 4000:
            SECURITY_INPUT_BLOCKED_TOTAL.labels(reason="too_long").inc()
        else:
            SECURITY_INPUT_BLOCKED_TOTAL.labels(reason="injection").inc()
        return ChatResponse(response=prompt_or_error)
    request = request.model_copy(update={"prompt": prompt_or_error})

    # Asegurarse de que el directorio de historiales exista
    os.makedirs(CHAT_HISTORIES_DIR, exist_ok=True)
    print(f"[DEBUG] ChatRequest: user_id='{request.user_id}', token_user='{user}' -> effective_user='{effective_user}'")
    
    # --- NUEVA LÓGICA PARA ORGANIZAR EL DÍA ---
    if "organiza mi día" in request.prompt.lower():
        try:
            headers = {"Authorization": f"Bearer {INTERNAL_API_SECRET}"}
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    f"{PRODUCTIVITY_SERVICE_URL}/organize",
                    params={"user_email": effective_user},
                    headers=headers,
                    timeout=60.0
                )
                response.raise_for_status() # Lanza un error si la petición falló (ej. 500)
                
                result_data = response.json()
                summary = result_data.get("summary", "Tu día ha sido organizado.")
                tasks_created = result_data.get("tasks_created", 0)
                final_response = f"{summary}\n\nHe creado {tasks_created} nuevas tareas en tu calendario."

            return ChatResponse(response=final_response)

        except httpx.RequestError as e:
            print(f"Error contacting productivity service: {e}")
            raise HTTPException(status_code=503, detail="El servicio de productividad no está disponible en este momento.")
        except httpx.HTTPStatusError as e:
            print(f"Productivity service returned an error: {e.response.text}")
            raise HTTPException(status_code=500, detail="Ocurrió un error al organizar tu día.")

    # --- LÓGICA DE CONFIRMACIÓN DE ALMACENAMIENTO ---
    affirmatives = ["si", "sí", "claro", "por favor", "guárdalo", "guardalo", "aceptar", "ok", "vale", "almacena", "guarda", "confirmo"]
    prompt_norm = request.prompt.lower().strip()
    if any(aff in prompt_norm for aff in affirmatives):
        pending_key = f"pending_doc:{effective_user}"
        pending_data_json = redis_client.get(pending_key)
        if pending_data_json:
            print(f"[DEBUG] Confirmación recibida para {pending_key}. Persistiendo a DB...")
            pending_data = json.loads(pending_data_json)
            save_user_document(
                user_id=pending_data["user_id"],
                doc_type=pending_data["doc_type"],
                summary=pending_data["summary"],
                content={}, 
                raw_analysis=pending_data["raw_analysis"]
            )
            redis_client.delete(pending_key)
            return ChatResponse(response="¡Listo! He guardado la información de tu ticket en la base de datos.")
        else:
            print(f"[DEBUG] No hay datos pendientes en {pending_key} para almacenar.")
    retrieval_keywords = ["mis tickets", "muestra mis tickets", "ver mis tickets", "buscar ticket"]
    if any(k in request.prompt.lower() for k in retrieval_keywords):
        docs = get_user_documents(effective_user)
        if not docs:
            return ChatResponse(response="No encontré ningún ticket guardado para ti.")
        
        response_text = "Aquí tienes tus últimos tickets:\n"
        for doc_type, summary, ts in docs:
            response_text += f"\n- {ts.strftime('%d/%m/%Y')}: {summary[:100]}..."
        return ChatResponse(response=response_text)


    # Determinar qué ID usar para el historial
    history_id = effective_user
    conv_id = request.conversation_id  # None → legacy flat history

    # --- MIGRACIÓN DE HISTORIAL (JSON -> DB) ---
    history = get_chat_history(history_id, conversation_id=conv_id)
    if not history:
        history_path = get_history_path_for_id(history_id)
        if os.path.exists(history_path):
            print(f"Migrando historial JSON para {history_id}...")
            with open(history_path, "r") as f:
                old_history = json.load(f)
                for msg in old_history:
                    save_chat_message(history_id, msg['role'], msg['content'])
            # Opcional: renombrar el archivo para no migrarlo de nuevo
            os.rename(history_path, history_path + ".migrated")
            history = get_chat_history(history_id)

    # 1. Recuperar documentos relevantes del USUARIO AUTENTICADO desde QDRANT
    # ELIMINADO: Se delega al Orquestador/Planner para evitar inyectar contexto irrelevante en saludos.
    context = ""
    # try:
    #     user_vectorstore = get_user_vectorstore(user)
    #     relevant_docs = user_vectorstore.similarity_search(request.prompt, k=3)
    #     if relevant_docs:
    #         RAG_HITS_TOTAL.inc()
    #     else:
    #         RAG_MISS_TOTAL.inc()
    #     context = "\n".join([doc.page_content for doc in relevant_docs])
    # except Exception as e:
    #     RAG_MISS_TOTAL.inc()
    #     print(f"[RAG/Qdrant] Error: {e}")
    #     context = ""

    # 2. Construir el historial de conversación
    conversation_history = ""
    if history:
        history_lines = [f"Human: {msg['content']}" if msg['role'] == 'user' else f"Assistant: {msg['content']}" for msg in history]
        conversation_history = "\n".join(history_lines)

    # 3. Crear el prompt final
    # Si hay imagen, usamos Vision Model con ChatOllama
    if request.image:
        try:
            print(f"[DEBUG] Imagen detectada de {history_id}. Procesando con {VISION_MODEL}...")
            # Instrucción de detección muy explícita
            detection_instruction = """
            INSTRUCCIÓN CRÍTICA: Si esta imagen es un ticket de compra, recibo o factura, debes:
            1. Analizarlo (fecha, lugar, total).
            2. Incluir EXPLICITAMENTE y SIN FALLAR la palabra clave [TICKET_DETECTED] al final de tu respuesta.
            ---
            """
            
            # Construir mensajes para ChatOllama
            content = [
                {"type": "text", "text": detection_instruction + request.prompt},
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{request.image}"},
                },
            ]
            
            # Invocación multimodal con instrucciones para detección de documentos
            system_content = "Eres Amael, un asistente experto en visión artificial. Si detectas un ticket, recibo o factura, DEBES terminar obligatoriamente con la etiqueta [TICKET_DETECTED]."
            if context:
                system_content += f"\n\nContexto relevante:\n{context}"

            messages = [
                SystemMessage(content=system_content),
                HumanMessage(content=content)
            ]
            
            print(f"[DEBUG] System Content: {system_content}")
            print(f"[DEBUG] User Prompt with Detection: {detection_instruction + request.prompt}")

            # Invocación multimodal
            start_vision_time = time.time()
            response = vision_llm.invoke(messages)
            vision_latency = time.time() - start_vision_time
            
            final_response = response.content
            print(f"[DEBUG] Vision LLM Latency: {vision_latency:.2f}s")
            print(f"[DEBUG] Raw Vision LLM Response: {final_response}")

            # Lógica de almacenamiento pendiente si se detecta ticket
            # Hacemos la comparación insensible a mayúsculas y con strip
            if "[TICKET_DETECTED]" in final_response.upper():
                print(f"[DEBUG] ¡Ticket detectado! Guardando en Redis para {history_id}...")
                # Limpiar el tag de la respuesta final para el usuario
                final_response = final_response.replace("[TICKET_DETECTED]", "").replace("[ticket_detected]", "").strip()
                final_response += "\n\n¿Te gustaría que guarde esta información de tu ticket?"
                
                # Guardar datos en Redis temporalmente (10 min) para esperar confirmación
                pending_data = {
                    "user_id": history_id,
                    "doc_type": "ticket",
                    "summary": final_response,
                    "raw_analysis": response.content 
                }
                redis_client.setex(f"pending_doc:{history_id}", 600, json.dumps(pending_data))
            else:
                print(f"[DEBUG] El modelo no marcó la imagen como ticket.")

            # Guardar en historial
            save_chat_message(history_id, "user", request.prompt + " [Imagen enviada]", conv_id)
            save_chat_message(history_id, "assistant", final_response, conv_id)

            return ChatResponse(response=final_response)

        except Exception as e:
            print(f"Error en Vision LLM: {e}")
            raise HTTPException(status_code=500, detail=f"Error al procesar la imagen con Qwen-VL: {e}")

    # --- NEW AGENTIC FLOW ---
    try:
        start_agent_time = time.time()
        
        # Tools wrapper for the executor
        def rag_tool(query: str):
            user_vectorstore = get_user_vectorstore(effective_user)
            relevant_docs = user_vectorstore.similarity_search(query, k=3)
            if relevant_docs:
                RAG_HITS_TOTAL.inc()
                return "\n".join([doc.page_content for doc in relevant_docs])
            RAG_MISS_TOTAL.inc()
            return "No relevant documents found."
            
        def productivity_tool(query: str):
            TOOL_CALLS_TOTAL.labels(tool="productivity").inc()
            headers = {"Authorization": f"Bearer {INTERNAL_API_SECRET}"}
            response = _http_client.post(
                f"{PRODUCTIVITY_SERVICE_URL}/organize",
                params={"user_email": effective_user},
                headers=headers,
                timeout=60.0
            )
            if response.status_code == 200:
                return response.json().get("summary", "Organized.")
            return "Error calling productivity service."

        def k8s_tool(query: str):
            # VALIDACIÓN DE WHITELIST DE INFRAESTRUCTURA
            if effective_user not in K8S_ALLOWED_USERS:
                return "Lo siento, tu usuario no cuenta con los privilegios de administrador requeridos para interactuar con la infraestructura del clúster."

            TOOL_CALLS_TOTAL.labels(tool="k8s").inc()
            K8S_AGENT_URL = "http://k8s-agent-service:8002"
            headers = {"Authorization": f"Bearer {INTERNAL_API_SECRET}"}
            response = _http_client.post(
                f"{K8S_AGENT_URL}/api/k8s-agent",
                json={"query": query, "user_email": effective_user},
                headers=headers,
                timeout=120.0
            )
            if response.status_code == 200:
                raw = response.json().get("response", "K8s info retrieved.")
                return f"```bash\n{raw}\n```"
            return "Error calling K8s agent."

        def web_search_tool(query: str) -> str:
            TOOL_CALLS_TOTAL.labels(tool="web_search").inc()
            return _web_search(query)

        def document_tool(query: str) -> str:
            TOOL_CALLS_TOTAL.labels(tool="document").inc()
            user_ctx = get_user_context(effective_user)
            import datetime as _dt
            today_str = _dt.date.today().strftime("%d de %B de %Y")
            doc_prompt = (
                f"Eres un redactor experto en documentos institucionales formales en español mexicano.\n"
                f"Fecha de hoy: {today_str}\n"
                f"{f'Perfil del autor: {user_ctx}' if user_ctx else ''}\n\n"
                f"Redacta el siguiente documento institucional: {query}\n\n"
                f"REGLAS:\n"
                f"1. Usa formato institucional formal con: encabezado, destinatario, cuerpo, conclusión y firma.\n"
                f"2. Usa markdown: # para título, ## para secciones, **negrita** para campos clave.\n"
                f"3. Incluye fecha, cargo del firmante y espacio para firma al final.\n"
                f"4. El documento debe estar listo para usar sin edición adicional.\n"
                f"5. Responde ÚNICAMENTE con el documento, sin comentarios previos ni posteriores."
            )
            content = _llm_doc.invoke(doc_prompt)
            return f"[DOCUMENT_START]\n{content}\n[DOCUMENT_END]"

        tools_map = {
            "rag": rag_tool,
            "productivity": productivity_tool,
            "k8s": k8s_tool,
            "web_search": web_search_tool,
            "document": document_tool,
        }

        # Restricción de solo lectura: sin K8S ni Productividad
        _readonly_prefix = ""
        if get_user_role(effective_user) == 'readonly':
            tools_map.pop("k8s", None)
            tools_map.pop("productivity", None)
            _readonly_prefix = "[RESTRICCIÓN ACTIVA: Este usuario es de solo lectura. NO uses K8S_TOOL ni PRODUCTIVITY_TOOL en ningún paso del plan.]\n\n"

        # P5-2: Reuse cached compiled orchestrator (compiled once at first request)
        orchestrator_app = get_orchestrator(redis_client=redis_client)

        initial_state = {
            "question": _readonly_prefix + request.prompt,
            "plan": [],
            "batches": [],
            "current_batch": 0,
            "current_step": 0,
            "context": context,
            "tool_results": [],
            "final_answer": None,
            "user_id": effective_user,
            # P3: Supervisor fields
            "retry_count": 0,
            "supervisor_score": 0,
            "supervisor_reason": "",
            # P5-2: tools injected through state so the compiled graph is reusable
            "tools_map": tools_map,
        }
        
        try:
            final_state = orchestrator_app.invoke(initial_state)
            
            agent_latency = time.time() - start_agent_time
            AGENT_EXECUTION_LATENCY.observe(agent_latency)
            PLANNER_STEPS_TOTAL.inc(len(final_state.get("plan", [])))
            
            logging.info(f"[AGENT] Success for user={effective_user} in {agent_latency:.2f}s")

            final_response = final_state.get("final_answer", "")
            if final_response is None:
                final_response = "Lo siento, no pude generar una respuesta en este momento."

            # P4-4: Sanitizar la salida antes de enviarla al usuario
            final_response = sanitize_output(final_response)

            # P3: log supervisor outcome
            sv_score = final_state.get("supervisor_score", 0)
            sv_decision = final_state.get("supervisor_decision", "N/A")
            sv_retries = final_state.get("retry_count", 0)
            logging.info(
                f"[SUPERVISOR] user={effective_user} decision={sv_decision} "
                f"score={sv_score} retries={sv_retries}"
            )

            # Save to history
            save_chat_message(history_id, "user", request.prompt, conv_id)
            save_chat_message(history_id, "assistant", final_response, conv_id)

            # Auto-title conversation from first user message
            if conv_id and len(history) == 0:
                title = request.prompt[:50].strip()
                if len(request.prompt) > 50:
                    title += "…"
                _touch_conversation(conv_id)
                pool = get_postgres_pool()
                if pool:
                    _c = pool.getconn()
                    try:
                        with _c.cursor() as cur:
                            cur.execute("UPDATE conversations SET title = %s WHERE id = %s AND title = 'Nueva conversación'", (title, conv_id))
                            _c.commit()
                    finally:
                        pool.putconn(_c)

            return ChatResponse(response=final_response)

        except Exception as orchestrator_exc:
            agent_latency = time.time() - start_agent_time
            AGENT_EXECUTION_LATENCY.observe(agent_latency)
            AGENT_FAILURES_TOTAL.inc()
            logging.error(f"[AGENT] Error in Agent Orchestrator for user={effective_user}: {orchestrator_exc}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"Error in agent orchestration: {orchestrator_exc}")

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"[AGENT] Unexpected outer error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error inesperado en el flujo del agente: {e}")

    
# --- ENDPOINTS DE CONOCIMIENTO SRE (Para k8s-agent) ---

SRE_COLLECTION = "amael_sre_knowledge"

def get_sre_vectorstore():
    """Obtiene el vectorstore específico para el conocimiento de SRE."""
    return QdrantVectorStore.from_existing_collection(
        embedding=_embeddings,
        collection_name=SRE_COLLECTION,
        url=QDRANT_URL,
    )

@app.post("/api/sre/ingest")
async def ingest_sre_knowledge(request: SREKnowledgeRequest, req: Request):
    """Permite indexar guías técnicas en la base de conocimiento de SRE."""
    # Validación de secreto interno para seguridad microservicio-microservicio
    token = req.headers.get("Authorization")
    if token != f"Bearer {INTERNAL_API_SECRET}":
         # Fallback para permitir ingesta manual si se desea, o restringir estrictamente
         # Por ahora permitimos si viene de localhost o con el secreto
         if req.client.host != "127.0.0.1":
            raise HTTPException(status_code=401, detail="No autorizado")

    try:
        if not _qdrant_client.collection_exists(SRE_COLLECTION):
            _qdrant_client.create_collection(
                collection_name=SRE_COLLECTION,
                vectors_config={"size": 768, "distance": "Cosine"}
            )

        vectorstore = QdrantVectorStore(
            client=_qdrant_client,
            collection_name=SRE_COLLECTION,
            embedding=_embeddings
        )

        from langchain_core.documents import Document
        doc = Document(page_content=request.content, metadata={"key": request.key})
        vectorstore.add_documents([doc])

        return {"message": f"Guía '{request.key}' indexada correctamente en SRE Knowledge."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/sre/query")
async def query_sre_knowledge(request: SREQueryRequest):
    """Consulta la base de conocimiento de SRE mediante búsqueda vectorial."""
    try:
        # Aquí no pedimos token para que el agente pueda consultar rápido,
        # pero restringimos la info que devuelve.
        if not _qdrant_client.collection_exists(SRE_COLLECTION):
            return {"response": "No hay base de conocimiento de SRE disponible aún."}

        vectorstore = QdrantVectorStore(
            client=_qdrant_client,
            collection_name=SRE_COLLECTION,
            embedding=_embeddings
        )
        
        results = vectorstore.similarity_search(request.query, k=2)
        if not results:
            return {"response": "No se encontró información relevante en el manual de SRE."}
            
        combined_info = "\n\n".join([doc.page_content for doc in results])
        return {"response": combined_info}
    except Exception as e:
        print(f"Error querying SRE RAG: {e}")
        return {"response": f"Error consultando conocimiento: {str(e)}"}

# ... TensorFlow (sin cambios, ya que no es multiusuario por ahora)
@app.post("/api/analyze-image", dependencies=[Depends(get_current_user)])
async def analyze_image(file: UploadFile = File(...), user: str = Depends(get_current_user)):
    """Endpoint para analizar una imagen con TensorFlow."""
    try:
        image_data = await file.read()
        img = Image.open(io.BytesIO(image_data)).convert('RGB')
        img = img.resize((224, 224))
        
        image_array = np.array(img, dtype=np.float32) / 255.0
        image_batch = np.expand_dims(image_array, axis=0)

        payload = {
            "instances": image_batch.tolist()
        }

        tf_serving_url = os.environ.get("TF_SERVING_URL")
        response = requests.post(tf_serving_url, json=payload)

        if response.status_code == 200:
            predictions = response.json()['predictions'][0]
            
            # --- NUEVA LÓGICA DE DECODIFICACIÓN ---
            # Obtenemos las etiquetas de ImageNet si no están cacheadas
            if not hasattr(app.state, "imagenet_labels"):
                url = "https://storage.googleapis.com/download.tensorflow.org/data/imagenet_class_index.json"
                try:
                    labels_response = requests.get(url, timeout=5)
                    app.state.imagenet_labels = labels_response.json()
                except Exception as e:
                    print(f"Error descargando etiquetas ImageNet: {e}")
                    return {"analysis_result": predictions[:5]} # Fallback
            
            labels_dict = app.state.imagenet_labels
            
            # Obtener los top 5 índices con mayor probabilidad
            top_indices = np.argsort(predictions)[-5:][::-1]
            
            results = []
            for idx in top_indices:
                class_id, label = labels_dict[str(idx)]
                prob = float(predictions[idx])
                results.append({
                    "etiqueta": label.replace("_", " "),
                    "probabilidad": f"{round(prob * 100, 2)}%"
                })
                
            resumen_ia = ", ".join([f"{item['etiqueta']} ({item['probabilidad']})" for item in results])
            
            return {
                "analisis_detallado": results,
                "resumen_ia": resumen_ia
            }
        else:
            raise HTTPException(status_code=500, detail="Error en TensorFlow Serving")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error procesando la imagen: {e}")

@app.post("/api/chat/stream", dependencies=[Depends(get_current_user)])
async def chat_stream(request: ChatRequest, user: str = Depends(get_current_user)):
    """Streaming SSE version of /api/chat. Emits status events then streams tokens."""
    from fastapi.responses import StreamingResponse as SR

    effective_user = resolve_user_id(request.user_id) if request.user_id else user

    # Rate limit
    allowed, _ = _check_rate_limit(effective_user)
    if not allowed:
        raise HTTPException(status_code=429, detail="Demasiadas solicitudes.")

    # Input validation
    is_valid, prompt_or_error = validate_prompt(request.prompt)
    if not is_valid:
        raise HTTPException(status_code=400, detail=prompt_or_error)

    prompt    = prompt_or_error
    conv_id   = request.conversation_id
    history_id = effective_user
    history   = get_chat_history(history_id, conversation_id=conv_id)

    conversation_history = ""
    if history:
        conversation_history = "\n".join(
            f"Human: {m['content']}" if m['role'] == 'user' else f"Assistant: {m['content']}"
            for m in history
        )

    # Memory Agent v1: cargar contexto del usuario
    user_memory_context = get_user_context(effective_user)

    # Restricción de solo lectura para stream endpoint
    _stream_is_readonly = get_user_role(effective_user) == 'readonly'
    _stream_readonly_prefix = (
        "[RESTRICCIÓN ACTIVA: Este usuario es de solo lectura. NO uses K8S_TOOL ni PRODUCTIVITY_TOOL en ningún paso del plan.]\n\n"
        if _stream_is_readonly else ""
    )

    def _run_agent_sync():
        """Synchronous agent execution — runs in thread pool."""
        def rag_tool(query: str):
            vs = get_user_vectorstore(effective_user)
            docs = vs.similarity_search(query, k=3)
            if docs:
                RAG_HITS_TOTAL.inc()
                return "\n".join(d.page_content for d in docs)
            RAG_MISS_TOTAL.inc()
            return "No relevant documents found."

        def productivity_tool(query: str):
            TOOL_CALLS_TOTAL.labels(tool="productivity").inc()
            r = _http_client.post(
                f"{PRODUCTIVITY_SERVICE_URL}/organize",
                params={"user_email": effective_user},
                headers={"Authorization": f"Bearer {INTERNAL_API_SECRET}"},
                timeout=60.0,
            )
            return r.json().get("summary", "Organized.") if r.status_code == 200 else "Error."

        def k8s_tool(query: str):
            if effective_user not in K8S_ALLOWED_USERS:
                return "Sin privilegios de administrador para esta operación."
            TOOL_CALLS_TOTAL.labels(tool="k8s").inc()
            r = _http_client.post(
                "http://k8s-agent-service:8002/api/k8s-agent",
                json={"query": query, "user_email": effective_user},
                headers={"Authorization": f"Bearer {INTERNAL_API_SECRET}"},
                timeout=120.0,
            )
            if r.status_code == 200:
                raw = r.json().get("response", "K8s info.")
                return f"```bash\n{raw}\n```"
            return "Error."

        def web_search_tool(query: str) -> str:
            TOOL_CALLS_TOTAL.labels(tool="web_search").inc()
            return _web_search(query)

        def document_tool(query: str) -> str:
            TOOL_CALLS_TOTAL.labels(tool="document").inc()
            user_ctx = get_user_context(effective_user)
            import datetime as _dt
            today_str = _dt.date.today().strftime("%d de %B de %Y")
            doc_prompt = (
                f"Eres un redactor experto en documentos institucionales formales en español mexicano.\n"
                f"Fecha de hoy: {today_str}\n"
                f"{f'Perfil del autor: {user_ctx}' if user_ctx else ''}\n\n"
                f"Redacta el siguiente documento institucional: {query}\n\n"
                f"REGLAS:\n"
                f"1. Usa formato institucional formal con: encabezado, destinatario, cuerpo, conclusión y firma.\n"
                f"2. Usa markdown: # para título, ## para secciones, **negrita** para campos clave.\n"
                f"3. Incluye fecha, cargo del firmante y espacio para firma al final.\n"
                f"4. El documento debe estar listo para usar sin edición adicional.\n"
                f"5. Responde ÚNICAMENTE con el documento, sin comentarios previos ni posteriores."
            )
            content = _llm_doc.invoke(doc_prompt)
            return f"[DOCUMENT_START]\n{content}\n[DOCUMENT_END]"

        _stream_tools_map = {
            "rag": rag_tool, "web_search": web_search_tool, "document": document_tool,
        }
        if not _stream_is_readonly:
            _stream_tools_map["k8s"] = k8s_tool
            _stream_tools_map["productivity"] = productivity_tool

        orch = get_orchestrator(redis_client=redis_client)
        initial_context = user_memory_context
        state = {
            "question": _stream_readonly_prefix + prompt,
            "plan": [], "batches": [], "current_batch": 0, "current_step": 0,
            "context": initial_context, "tool_results": [], "final_answer": None,
            "user_id": effective_user, "retry_count": 0,
            "supervisor_score": 0, "supervisor_reason": "",
            "tools_map": _stream_tools_map,
        }
        final_state = orch.invoke(state)
        answer = final_state.get("final_answer") or "Lo siento, no pude generar una respuesta."
        return sanitize_output(answer), final_state

    def _sse(event_type: str, payload: dict) -> str:
        return f"data: {json.dumps({'type': event_type, **payload})}\n\n"

    async def generate():
        yield _sse("status", {"msg": "🧠 Analizando consulta…"})

        # Run orchestrator in thread pool; poll with status updates
        loop = asyncio.get_event_loop()
        future = loop.run_in_executor(None, _run_agent_sync)

        status_schedule = [
            (2.0,  "📋 Creando plan de ejecución…"),
            (5.0,  "⚡ Ejecutando herramientas…"),
            (10.0, "✍️ Sintetizando respuesta…"),
            (20.0, "⏳ Procesando consulta compleja…"),
        ]
        elapsed = 0.0
        poll    = 0.25
        si      = 0

        while not future.done():
            await asyncio.sleep(poll)
            elapsed += poll
            if si < len(status_schedule) and elapsed >= status_schedule[si][0]:
                yield _sse("status", {"msg": status_schedule[si][1]})
                si += 1

        try:
            final_answer, final_state = future.result()
        except Exception as e:
            yield _sse("error", {"msg": str(e)})
            return

        # Persist
        save_chat_message(history_id, "user",      prompt,       conv_id)
        save_chat_message(history_id, "assistant", final_answer, conv_id)
        if conv_id and len(history) == 0:
            title = prompt[:50].strip() + ("…" if len(prompt) > 50 else "")
            _touch_conversation(conv_id)
            pool = get_postgres_pool()
            if pool:
                _c = pool.getconn()
                try:
                    with _c.cursor() as cur:
                        cur.execute(
                            "UPDATE conversations SET title=%s WHERE id=%s AND title='Nueva conversación'",
                            (title, conv_id),
                        )
                        _c.commit()
                finally:
                    pool.putconn(_c)
        
        # Memory Agent v1: extraer hechos en background (no bloquea el stream)
        asyncio.create_task(extract_facts_background(effective_user, prompt, final_answer, conv_id))

        # Stream tokens word-by-word
        words = final_answer.split(" ")
        for i, word in enumerate(words):
            chunk = word + (" " if i < len(words) - 1 else "")
            yield _sse("token", {"content": chunk})
            await asyncio.sleep(0.018)

        yield _sse("done", {"conv_id": conv_id})

    return SR(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.get("/api/conversations", dependencies=[Depends(get_current_user)])
async def list_conversations(user: str = Depends(get_current_user)):
    return {"conversations": get_user_conversations(user)}


@app.post("/api/conversations", dependencies=[Depends(get_current_user)])
async def new_conversation(body: ConversationCreate, user: str = Depends(get_current_user)):
    conv_id = create_conversation(user, body.title)
    if conv_id is None:
        raise HTTPException(status_code=503, detail="Base de datos no disponible")
    return {"id": conv_id, "title": body.title}


@app.get("/api/conversations/{conv_id}/messages", dependencies=[Depends(get_current_user)])
async def conversation_messages(conv_id: int, user: str = Depends(get_current_user)):
    return {"messages": get_conversation_messages(conv_id, user)}


class ConversationUpdateRequest(BaseModel):
    title: str

@app.patch("/api/conversations/{conv_id}", dependencies=[Depends(get_current_user)])
async def update_conversation(conv_id: int, body: ConversationUpdateRequest, user: str = Depends(get_current_user)):
    pool = get_postgres_pool()
    if not pool:
        return {"ok": True}
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE conversations SET title=%s WHERE id=%s AND user_id=%s",
                (body.title[:100], conv_id, user)
            )
        conn.commit()
        return {"ok": True}
    finally:
        pool.putconn(conn)


@app.delete("/api/conversations/{conv_id}", dependencies=[Depends(get_current_user)])
async def delete_conversation(conv_id: int, user: str = Depends(get_current_user)):
    pool = get_postgres_pool()
    if not pool:
        return {"ok": True}
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM chat_history WHERE conversation_id=%s AND user_id=%s", (conv_id, user))
            cur.execute("DELETE FROM message_feedback WHERE conversation_id=%s AND user_id=%s", (conv_id, user))
            cur.execute("DELETE FROM conversations WHERE id=%s AND user_id=%s", (conv_id, user))
        conn.commit()
        return {"ok": True}
    finally:
        pool.putconn(conn)


@app.post("/api/feedback", dependencies=[Depends(get_current_user)])
async def save_feedback(body: FeedbackRequest, user: str = Depends(get_current_user)):
    pool = get_postgres_pool()
    if not pool:
        return {"ok": True}
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO message_feedback (user_id, conversation_id, message_index, sentiment) VALUES (%s, %s, %s, %s)",
                (user, body.conversation_id, body.message_index, body.sentiment)
            )
            conn.commit()
    finally:
        pool.putconn(conn)
    return {"ok": True}


# ═══════════════════════════════════════════════════════════════════════════════
# MEMORY AGENT v1 — Perfil + Hechos del usuario
# ═══════════════════════════════════════════════════════════════════════════════

def get_user_context(user_id: str) -> str:
    """Carga el perfil y hechos recientes del usuario para inyectar en el agente."""
    pool = get_postgres_pool()
    if not pool:
        return ""
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            # Perfil básico
            cur.execute("SELECT display_name, preferences, context_data FROM user_profile WHERE user_id=%s", (user_id,))
            row = cur.fetchone()
            profile_lines = []
            if row:
                name, prefs, ctx = row
                if name:
                    profile_lines.append(f"Nombre del usuario: {name}")
                if prefs:
                    for k, v in prefs.items():
                        profile_lines.append(f"Preferencia — {k}: {v}")
                if ctx:
                    for k, v in ctx.items():
                        profile_lines.append(f"Contexto — {k}: {v}")

            # Hechos recientes (últimos 15)
            cur.execute(
                "SELECT fact, category FROM user_facts WHERE user_id=%s ORDER BY created_at DESC LIMIT 15",
                (user_id,)
            )
            facts = cur.fetchall()
            fact_lines = [f"- [{cat}] {fact}" for fact, cat in facts] if facts else []

        if not profile_lines and not fact_lines:
            return ""

        parts = []
        if profile_lines:
            parts.append("Perfil del usuario:\n" + "\n".join(profile_lines))
        if fact_lines:
            parts.append("Hechos conocidos sobre el usuario:\n" + "\n".join(fact_lines))
        return "\n\n".join(parts)
    except Exception as e:
        logging.warning(f"[MEMORY] Error cargando contexto de {user_id}: {e}")
        return ""
    finally:
        pool.putconn(conn)


def save_user_fact(user_id: str, fact: str, category: str = "general", conv_id: Optional[int] = None):
    """Persiste un hecho extraído sobre el usuario."""
    pool = get_postgres_pool()
    if not pool or not fact.strip():
        return
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            # Evitar duplicados exactos
            cur.execute(
                "SELECT id FROM user_facts WHERE user_id=%s AND fact=%s",
                (user_id, fact.strip())
            )
            if cur.fetchone():
                return
            cur.execute(
                "INSERT INTO user_facts (user_id, fact, category, source_conv_id) VALUES (%s, %s, %s, %s)",
                (user_id, fact.strip()[:500], category, conv_id)
            )
        conn.commit()
    except Exception as e:
        logging.warning(f"[MEMORY] Error guardando hecho: {e}")
    finally:
        pool.putconn(conn)


async def extract_facts_background(user_id: str, question: str, answer: str, conv_id: Optional[int] = None):
    """Extrae hechos relevantes de la conversación y los guarda. Falla silenciosamente."""
    try:
        extract_prompt = f"""Analiza este intercambio y extrae hechos *duraderos* sobre el usuario (máx 3).
Solo extrae: preferencias, proyectos activos, objetivos, contexto personal/laboral, herramientas que usa.
NO extraigas: preguntas puntuales, respuestas técnicas temporales, saludos.

Usuario preguntó: {question[:400]}
Respuesta del asistente: {answer[:600]}

Responde ÚNICAMENTE con JSON válido:
{{"facts": [{{"text": "hecho 1", "category": "preferencia|proyecto|objetivo|contexto"}}, ...]}}
Si no hay hechos relevantes: {{"facts": []}}"""

        from langchain_ollama import ChatOllama as _CO
        from langchain_core.messages import HumanMessage as _HM
        _llm = _CO(model=LLM_MODEL, base_url=OLLAMA_BASE_URL, temperature=0, num_predict=200)
        response = await asyncio.to_thread(_llm.invoke, [_HM(content=extract_prompt)])
        raw = response.content.strip()
        # Extraer JSON aunque tenga texto alrededor
        import re as _re
        m = _re.search(r'\{.*\}', raw, _re.DOTALL)
        if m:
            data = json.loads(m.group())
            for item in data.get("facts", [])[:3]:
                save_user_fact(user_id, item.get("text", ""), item.get("category", "general"), conv_id)
    except Exception as e:
        logging.debug(f"[MEMORY] extract_facts_background error (non-critical): {e}")


def upsert_user_profile(user_id: str, display_name: Optional[str] = None,
                         preferences: Optional[dict] = None, context_data: Optional[dict] = None):
    """Crea o actualiza el perfil del usuario."""
    pool = get_postgres_pool()
    if not pool:
        return
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT user_id FROM user_profile WHERE user_id=%s", (user_id,))
            if cur.fetchone():
                updates, vals = [], []
                if display_name:
                    updates.append("display_name=%s"); vals.append(display_name)
                if preferences:
                    updates.append("preferences=preferences || %s::jsonb"); vals.append(json.dumps(preferences))
                if context_data:
                    updates.append("context_data=context_data || %s::jsonb"); vals.append(json.dumps(context_data))
                if updates:
                    updates.append("updated_at=NOW()")
                    vals.append(user_id)
                    cur.execute(f"UPDATE user_profile SET {', '.join(updates)} WHERE user_id=%s", vals)
            else:
                cur.execute(
                    "INSERT INTO user_profile (user_id, display_name, preferences, context_data) VALUES (%s,%s,%s,%s)",
                    (user_id, display_name, json.dumps(preferences or {}), json.dumps(context_data or {}))
                )
        conn.commit()
    except Exception as e:
        logging.warning(f"[MEMORY] upsert_user_profile error: {e}")
    finally:
        pool.putconn(conn)


# ═══════════════════════════════════════════════════════════════════════════════
# WHATSAPP HELPER — Envío proactivo de mensajes
# ═══════════════════════════════════════════════════════════════════════════════

async def send_whatsapp_message(phone: str, text: str) -> bool:
    """Envía un mensaje de texto por WhatsApp via el bridge. Retorna True si tuvo éxito."""
    try:
        r = await asyncio.to_thread(
            _http_client.post,
            f"{WHATSAPP_BRIDGE_URL}/send",
            json={"phoneNumber": phone, "text": text},
            timeout=15.0,
        )
        return r.status_code == 200
    except Exception as e:
        logging.warning(f"[WA] Error enviando mensaje a {phone}: {e}")
        return False


# ═══════════════════════════════════════════════════════════════════════════════
# TAREA #2 — Reporte de calidad del supervisor
# ═══════════════════════════════════════════════════════════════════════════════

@app.get("/api/admin/quality-report")
async def quality_report(user: str = Depends(get_current_user)):
    """Analiza los scores del supervisor y devuelve estadísticas de calidad."""
    if user not in K8S_ALLOWED_USERS:
        raise HTTPException(status_code=403, detail="Solo administradores.")

    all_entries = []
    try:
        keys = redis_client.keys("agent_feedback:*")
        for key in keys:
            entries = redis_client.lrange(key, 0, -1)
            for e in entries:
                try:
                    all_entries.append(json.loads(e))
                except Exception:
                    pass
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Redis no disponible: {e}")

    if not all_entries:
        return {"total": 0, "message": "No hay datos de feedback aún."}

    scores   = [e.get("score", 0) for e in all_entries if "score" in e]
    decisions = [e.get("decision", "") for e in all_entries]
    replans  = decisions.count("REPLAN")
    accepts  = decisions.count("ACCEPT")

    # Distribución de scores
    dist = {str(i): 0 for i in range(11)}
    for s in scores:
        dist[str(min(int(s), 10))] += 1

    # Sesiones con score bajo (< 6)
    low_sessions = [e for e in all_entries if e.get("score", 10) < 6]
    low_questions = [e.get("question", "")[:100] for e in low_sessions[:10]]

    avg = round(sum(scores) / len(scores), 2) if scores else 0

    return {
        "total_evaluaciones": len(all_entries),
        "score_promedio": avg,
        "score_min": min(scores) if scores else 0,
        "score_max": max(scores) if scores else 0,
        "distribución_scores": dist,
        "total_accept": accepts,
        "total_replan": replans,
        "tasa_replan_pct": round(replans / len(decisions) * 100, 1) if decisions else 0,
        "sesiones_baja_calidad": len(low_sessions),
        "ejemplos_baja_calidad": low_questions,
    }


# ═══════════════════════════════════════════════════════════════════════════════
# TAREA #4 — Memory Agent: endpoints de perfil y hechos
# ═══════════════════════════════════════════════════════════════════════════════

class UserProfileUpdate(BaseModel):
    display_name: Optional[str] = None
    timezone: Optional[str] = None
    preferences: Optional[dict] = None
    context_data: Optional[dict] = None

@app.get("/api/memory/profile", dependencies=[Depends(get_current_user)])
async def get_profile(user: str = Depends(get_current_user)):
    """Devuelve el perfil y hechos del usuario autenticado."""
    context = get_user_context(user)
    pool = get_postgres_pool()
    if not pool:
        return {"profile": None, "facts": [], "context_text": ""}
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT display_name, preferences, context_data, updated_at, role, timezone FROM user_profile WHERE user_id=%s", (user,))
            row = cur.fetchone()
            profile = {"display_name": row[0], "preferences": row[1], "context_data": row[2], "updated_at": str(row[3]), "role": row[4] or 'user', "timezone": row[5]} if row else None
            cur.execute("SELECT fact, category, created_at FROM user_facts WHERE user_id=%s ORDER BY created_at DESC LIMIT 50", (user,))
            facts = [{"fact": r[0], "category": r[1], "date": str(r[2])} for r in cur.fetchall()]
        return {"profile": profile, "facts": facts, "context_text": context}
    finally:
        pool.putconn(conn)


@app.patch("/api/memory/profile", dependencies=[Depends(get_current_user)])
async def update_profile(body: UserProfileUpdate, user: str = Depends(get_current_user)):
    """Actualiza el perfil del usuario."""
    upsert_user_profile(user, body.display_name, body.preferences, body.context_data)
    pool = get_postgres_pool()
    if body.timezone and pool:
        conn = pool.getconn()
        try:
            with conn.cursor() as cur:
                cur.execute("UPDATE user_profile SET timezone = %s WHERE user_id = %s", (body.timezone, user))
                conn.commit()
        finally:
            pool.putconn(conn)
    return {"ok": True}


@app.get("/api/memory/goals", dependencies=[Depends(get_current_user)])
async def list_goals(user: str = Depends(get_current_user)):
    """Lista los objetivos activos del usuario."""
    pool = get_postgres_pool()
    if not pool:
        return {"goals": []}
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, title, category, status, progress, deadline, milestones, created_at FROM user_goals WHERE user_id=%s AND status='active' ORDER BY created_at DESC",
                (user,)
            )
            goals = [
                {"id": r[0], "title": r[1], "category": r[2], "status": r[3],
                 "progress": r[4], "deadline": str(r[5]) if r[5] else None,
                 "milestones": r[6], "created_at": str(r[7])}
                for r in cur.fetchall()
            ]
        return {"goals": goals}
    finally:
        pool.putconn(conn)


class GoalCreate(BaseModel):
    title: str
    category: str = "personal"
    deadline: Optional[str] = None

@app.post("/api/memory/goals", dependencies=[Depends(get_current_user)])
async def create_goal(body: GoalCreate, user: str = Depends(get_current_user)):
    """Crea un nuevo objetivo."""
    pool = get_postgres_pool()
    if not pool:
        raise HTTPException(status_code=503, detail="DB no disponible")
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO user_goals (user_id, title, category, deadline) VALUES (%s,%s,%s,%s) RETURNING id",
                (user, body.title[:200], body.category, body.deadline)
            )
            goal_id = cur.fetchone()[0]
        conn.commit()
        return {"id": goal_id, "title": body.title}
    finally:
        pool.putconn(conn)


# ═══════════════════════════════════════════════════════════════════════════════
# TAREA #5 — Day Planner matutino
# ═══════════════════════════════════════════════════════════════════════════════

async def _get_daily_news() -> str:
    """Obtiene 3 noticias relevantes del día sobre AI, infraestructura y gobierno digital."""
    topics = [
        "inteligencia artificial gobierno digital Mexico 2025",
        "Kubernetes infraestructura cloud noticias",
        "transformacion digital instituciones publicas",
    ]
    news_parts = []
    for topic in topics:
        try:
            result = await asyncio.to_thread(_web_search, topic)
            # Take first 300 chars of the result as a headline summary
            first_line = result.strip().split('\n')[0][:280]
            if first_line and "Error" not in first_line:
                news_parts.append(f"• {first_line}")
        except Exception:
            pass
    return "\n".join(news_parts) if news_parts else ""


async def _get_k8s_health_brief(user_id: str) -> str:
    """Obtiene un resumen breve del estado del cluster via k8s-agent."""
    try:
        r = await asyncio.to_thread(
            _http_client.post,
            f"{K8S_AGENT_URL}/api/k8s-agent",
            json={"query": "Dame un resumen breve del estado del cluster: pods con problemas, uso de recursos.", "user_email": user_id},
            headers={"Authorization": f"Bearer {INTERNAL_API_SECRET}"},
            timeout=60.0,
        )
        return r.json().get("response", "") if r.status_code == 200 else ""
    except Exception as e:
        logging.warning(f"[PLANNER] Error k8s health: {e}")
        return ""


async def _get_calendar_brief(user_id: str) -> str:
    """Obtiene eventos de hoy via productivity-service."""
    try:
        r = await asyncio.to_thread(
            _http_client.post,
            f"{PRODUCTIVITY_SERVICE_URL}/organize",
            headers={"Authorization": f"Bearer {INTERNAL_API_SECRET}"},
            params={"user_email": user_id},
            timeout=60.0,
        )
        data = r.json() if r.status_code == 200 else {}
        return data.get("summary", "") if data else ""
    except Exception as e:
        logging.warning(f"[PLANNER] Error calendar: {e}")
        return ""


async def _get_active_goals_text(user_id: str) -> str:
    """Obtiene objetivos activos como texto."""
    pool = get_postgres_pool()
    if not pool:
        return ""
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT title, progress, category FROM user_goals WHERE user_id=%s AND status='active' ORDER BY updated_at DESC LIMIT 5",
                (user_id,)
            )
            rows = cur.fetchall()
        if not rows:
            return ""
        return "\n".join([f"- {r[0]} ({r[2]}, {r[1]}% completado)" for r in rows])
    except Exception:
        return ""
    finally:
        pool.putconn(conn)


# ═══════════════════════════════════════════════════════════════════════════════
# ADMIN — Gestión de usuarios, identidades y configuración de plataforma
# ═══════════════════════════════════════════════════════════════════════════════

class AdminUserCreate(BaseModel):
    email: str
    role: str = 'user'
    phone: Optional[str] = None
    display_name: Optional[str] = None

class AdminUserUpdate(BaseModel):
    role: Optional[str] = None
    status: Optional[str] = None
    display_name: Optional[str] = None

class IdentityCreate(BaseModel):
    identity_type: str   # 'whatsapp' | 'email' | 'telegram'
    identity_value: str

class AccessRequestCreate(BaseModel):
    phone: str
    name: Optional[str] = None

def _require_admin(user: str):
    if get_user_role(user) != 'admin':
        raise HTTPException(status_code=403, detail="Solo administradores.")


@app.get("/api/admin/users", dependencies=[Depends(get_current_user)])
async def admin_list_users(user: str = Depends(get_current_user)):
    _require_admin(user)
    pool = get_postgres_pool()
    if not pool:
        return {"users": []}
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT p.user_id, p.display_name, p.role, p.status,
                       array_agg(i.identity_value || '|' || i.identity_type)
                           FILTER (WHERE i.id IS NOT NULL AND i.identity_type != 'email') as identities
                FROM user_profile p
                LEFT JOIN user_identities i ON i.canonical_user_id = p.user_id
                GROUP BY p.user_id, p.display_name, p.role, p.status
                ORDER BY CASE p.role WHEN 'admin' THEN 0 ELSE 1 END, p.user_id
            """)
            users = []
            for r in cur.fetchall():
                identities = []
                if r[4]:
                    for id_str in r[4]:
                        val, typ = id_str.rsplit('|', 1)
                        identities.append({"type": typ, "value": val})
                users.append({
                    "user_id": r[0], "display_name": r[1],
                    "role": r[2] or 'user', "status": r[3] or 'active',
                    "identities": identities,
                })
            return {"users": users}
    finally:
        pool.putconn(conn)


@app.post("/api/admin/users", dependencies=[Depends(get_current_user)])
async def admin_create_user(body: AdminUserCreate, user: str = Depends(get_current_user)):
    _require_admin(user)
    pool = get_postgres_pool()
    if not pool:
        raise HTTPException(status_code=503)
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO user_profile (user_id, display_name, role, status)
                VALUES (%s, %s, %s, 'active')
                ON CONFLICT (user_id) DO UPDATE
                SET role = EXCLUDED.role, status = 'active',
                    display_name = COALESCE(EXCLUDED.display_name, user_profile.display_name)
            """, (body.email, body.display_name, body.role))
            cur.execute("""
                INSERT INTO user_identities (canonical_user_id, identity_type, identity_value)
                VALUES (%s, 'email', %s) ON CONFLICT (identity_value) DO NOTHING
            """, (body.email, body.email))
            if body.phone:
                cur.execute("""
                    INSERT INTO user_identities (canonical_user_id, identity_type, identity_value)
                    VALUES (%s, 'whatsapp', %s)
                    ON CONFLICT (identity_value) DO UPDATE SET canonical_user_id = EXCLUDED.canonical_user_id
                """, (body.email, body.phone))
            conn.commit()
        return {"ok": True, "user_id": body.email}
    finally:
        pool.putconn(conn)


@app.patch("/api/admin/users/{target_uid:path}", dependencies=[Depends(get_current_user)])
async def admin_update_user(target_uid: str, body: AdminUserUpdate, user: str = Depends(get_current_user)):
    _require_admin(user)
    pool = get_postgres_pool()
    if not pool:
        raise HTTPException(status_code=503)
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            if body.role:
                cur.execute("UPDATE user_profile SET role = %s WHERE user_id = %s", (body.role, target_uid))
            if body.status:
                cur.execute("UPDATE user_profile SET status = %s WHERE user_id = %s", (body.status, target_uid))
            if body.display_name is not None:
                cur.execute("UPDATE user_profile SET display_name = %s WHERE user_id = %s", (body.display_name, target_uid))
            conn.commit()
        return {"ok": True}
    finally:
        pool.putconn(conn)


@app.delete("/api/admin/users/{target_uid:path}", dependencies=[Depends(get_current_user)])
async def admin_delete_user(target_uid: str, user: str = Depends(get_current_user)):
    """Elimina permanentemente un usuario (hard delete)."""
    _require_admin(user)
    if target_uid == user:
        raise HTTPException(status_code=400, detail="No puedes eliminarte a ti mismo.")
    pool = get_postgres_pool()
    if not pool:
        raise HTTPException(status_code=503)
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM user_identities WHERE canonical_user_id = %s", (target_uid,))
            cur.execute("DELETE FROM user_profile WHERE user_id = %s", (target_uid,))
            conn.commit()
        return {"ok": True}
    finally:
        pool.putconn(conn)


@app.post("/api/admin/users/{target_uid:path}/identity", dependencies=[Depends(get_current_user)])
async def admin_add_identity(target_uid: str, body: IdentityCreate, user: str = Depends(get_current_user)):
    _require_admin(user)
    pool = get_postgres_pool()
    if not pool:
        raise HTTPException(status_code=503)
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO user_identities (canonical_user_id, identity_type, identity_value)
                VALUES (%s, %s, %s)
                ON CONFLICT (identity_value) DO UPDATE SET canonical_user_id = EXCLUDED.canonical_user_id
            """, (target_uid, body.identity_type, body.identity_value))
            conn.commit()
        return {"ok": True}
    finally:
        pool.putconn(conn)


@app.delete("/api/admin/users/{target_uid:path}/identity/{identity_value}", dependencies=[Depends(get_current_user)])
async def admin_remove_identity(target_uid: str, identity_value: str, user: str = Depends(get_current_user)):
    _require_admin(user)
    pool = get_postgres_pool()
    if not pool:
        raise HTTPException(status_code=503)
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "DELETE FROM user_identities WHERE canonical_user_id = %s AND identity_value = %s AND identity_type != 'email'",
                (target_uid, identity_value)
            )
            conn.commit()
        return {"ok": True}
    finally:
        pool.putconn(conn)


@app.get("/api/admin/settings", dependencies=[Depends(get_current_user)])
async def admin_get_settings(user: str = Depends(get_current_user)):
    _require_admin(user)
    return {
        "allow_access_requests": get_platform_setting("allow_access_requests", "false") == "true",
    }


@app.patch("/api/admin/settings", dependencies=[Depends(get_current_user)])
async def admin_update_settings(body: dict, user: str = Depends(get_current_user)):
    _require_admin(user)
    pool = get_postgres_pool()
    if not pool:
        raise HTTPException(status_code=503)
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            for key, value in body.items():
                cur.execute("""
                    INSERT INTO platform_settings (key, value) VALUES (%s, %s)
                    ON CONFLICT (key) DO UPDATE SET value = EXCLUDED.value, updated_at = NOW()
                """, (key, str(value).lower()))
            conn.commit()
        return {"ok": True}
    finally:
        pool.putconn(conn)


@app.get("/api/identity/check")
async def identity_check(identifier: str, request: Request):
    """Endpoint interno para que el WhatsApp bridge verifique acceso y resuelva usuario canónico."""
    auth = request.headers.get("Authorization", "")
    if auth.removeprefix("Bearer ").strip() != INTERNAL_API_SECRET:
        raise HTTPException(status_code=403)
    allowed, canonical = is_user_registered(identifier)
    allow_requests = get_platform_setting("allow_access_requests", "false") == "true"
    return {"allowed": allowed, "canonical_user_id": canonical, "allow_requests": allow_requests}


@app.post("/api/auth/access-request")
async def access_request(body: AccessRequestCreate):
    """Solicitud de acceso de usuario no registrado."""
    if get_platform_setting("allow_access_requests", "false") != "true":
        raise HTTPException(status_code=403, detail="Las solicitudes de acceso están desactivadas.")
    admin_phone = os.environ.get('ADMIN_PHONE', '5219993437008')
    msg = f"🔔 *Solicitud de acceso a Amael-IA*\n\nNúmero: +{body.phone}"
    if body.name:
        msg += f"\nNombre: {body.name}"
    msg += "\n\nAproba desde el Panel Admin en la app."
    asyncio.create_task(send_whatsapp_message(admin_phone, msg))
    return {"ok": True}


# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/api/planner/daily")
async def daily_planner(request: Request):
    """
    Genera y envía el plan del día por WhatsApp.
    Llamado por el CronJob de Kubernetes cada mañana.
    Autenticado via INTERNAL_API_SECRET en el header.
    """
    auth = request.headers.get("Authorization", "")
    if auth.removeprefix("Bearer ").strip() != INTERNAL_API_SECRET:
        raise HTTPException(status_code=403, detail="No autorizado.")

    body = await request.json()
    user_id = body.get("user_id", ADMIN_PHONE)
    phone   = body.get("phone",   ADMIN_PHONE)

    import datetime as _dt
    _days_es   = ["lunes","martes","miércoles","jueves","viernes","sábado","domingo"]
    _months_es = ["enero","febrero","marzo","abril","mayo","junio",
                  "julio","agosto","septiembre","octubre","noviembre","diciembre"]
    _now  = _dt.date.today()
    today = f"{_days_es[_now.weekday()]} {_now.day} de {_months_es[_now.month-1]} de {_now.year}"

    # Recopilar todo el contexto en paralelo
    k8s_health, calendar_summary, goals_text, user_ctx, news_text = await asyncio.gather(
        _get_k8s_health_brief(user_id),
        _get_calendar_brief(user_id),
        _get_active_goals_text(user_id),
        asyncio.to_thread(get_user_context, user_id),
        _get_daily_news(),
    )

    plan_prompt = f"""Eres el asistente ejecutivo personal de un Subdirector de Infraestructura Digital.
Genera su briefing ejecutivo matutino. Hoy es {today}.

{f"Perfil:{chr(10)}{user_ctx}" if user_ctx else ""}

{f"📅 AGENDA DE HOY:{chr(10)}{calendar_summary}" if calendar_summary else "📅 Sin eventos registrados en el calendario."}

{f"🖥️ ESTADO DEL CLÚSTER:{chr(10)}{k8s_health}" if k8s_health else "🖥️ Clúster: Sin datos disponibles."}

{f"🎯 OBJETIVOS ACTIVOS:{chr(10)}{goals_text}" if goals_text else "🎯 Sin objetivos activos registrados."}

{f"📰 NOTICIAS DEL DÍA:{chr(10)}{news_text}" if news_text else ""}

REGLA CRÍTICA: USA ÚNICAMENTE la información proporcionada arriba. NUNCA inventes eventos, reuniones, citas ni datos que no estén en el contexto. Si una sección dice "Sin eventos" o "Sin datos", repórtalo tal cual.

Genera el briefing ejecutivo con este formato WhatsApp (*negritas*, emojis, saltos de línea):
1. 🌅 Saludo ejecutivo personalizado con el día y fecha
2. 📅 Agenda: si hay eventos en AGENDA DE HOY, lista máximo 3; si dice "Sin eventos", escribe "📅 *Agenda limpia* — sin reuniones programadas hoy."
3. 🖥️ Infraestructura: estado real del clúster según ESTADO DEL CLÚSTER (✅ OK o ⚠️ alertas)
4. 🎯 Prioridad del día: solo si hay objetivos en OBJETIVOS ACTIVOS; si no hay, omite esta sección
5. 📰 Una noticia del sector si está disponible en NOTICIAS DEL DÍA
6. 💡 Una pregunta de reflexión estratégica para el día

Tono: ejecutivo, conciso, orientado a la acción. Máximo 350 palabras."""

    try:
        from langchain_ollama import ChatOllama as _CO
        from langchain_core.messages import HumanMessage as _HM
        _llm = _CO(model=LLM_MODEL, base_url=OLLAMA_BASE_URL, temperature=0.4, num_predict=600)
        response = await asyncio.to_thread(_llm.invoke, [_HM(content=plan_prompt)])
        plan_text = response.content.strip()
    except Exception as e:
        logging.error(f"[PLANNER] Error generando plan: {e}")
        plan_text = f"🌅 Buenos días! Son las {_dt.datetime.now().strftime('%H:%M')}. No pude generar el plan completo hoy, pero estoy disponible para ayudarte."

    # Enviar por WhatsApp
    sent = await send_whatsapp_message(phone, plan_text)
    logging.info(f"[PLANNER] Plan enviado a {phone}: {sent}")

    return {"ok": True, "sent": sent, "plan": plan_text}


# ═══════════════════════════════════════════════════════════════════════════════
# WHATSAPP PERSONAL — Gestión de sesión y settings
# ═══════════════════════════════════════════════════════════════════════════════

class WhatsappPersonalSettings(BaseModel):
    auto_reply:        Optional[bool] = None
    quiet_enabled:     Optional[bool] = None   # activa/desactiva el horario silencioso
    quiet_start:       Optional[int]  = None   # hora 0-23
    quiet_end:         Optional[int]  = None   # hora 0-23
    active_days:       Optional[list] = None   # [1,2,3,4,5] = Lun-Vie (ISO: 1=Lun, 7=Dom)
    reply_scope:       Optional[str]  = None   # 'all' | 'contacts_only' | 'no_groups' | 'custom'
    allowed_contacts:  Optional[list] = None   # ["5219993437008", ...] solo para scope=custom
    ai_assist:         Optional[bool] = None
    offline_msg:       Optional[str]  = None


def _get_wp_settings(user_id: str) -> dict:
    """Lee settings de whatsapp_personal para un usuario."""
    defaults = {
        "auto_reply": True, "quiet_enabled": True, "quiet_start": 22, "quiet_end": 8,
        "active_days": [1, 2, 3, 4, 5], "reply_scope": "all", "allowed_contacts": [],
        "ai_assist": True, "offline_msg": None,
    }
    pool = get_postgres_pool()
    if not pool:
        return defaults
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT auto_reply, quiet_enabled, quiet_start, quiet_end, active_days, "
                "reply_scope, allowed_contacts, ai_assist, offline_msg "
                "FROM whatsapp_personal_settings WHERE user_id=%s", (user_id,)
            )
            row = cur.fetchone()
            if not row:
                return defaults
            return {
                "auto_reply":       row[0],
                "quiet_enabled":    row[1] if row[1] is not None else True,
                "quiet_start":      row[2], "quiet_end":       row[3],
                "active_days":      list(row[4]) if row[4] else defaults["active_days"],
                "reply_scope":      row[5],
                "allowed_contacts": list(row[6]) if row[6] else [],
                "ai_assist":        row[7], "offline_msg":     row[8],
            }
    finally:
        pool.putconn(conn)


def _is_in_quiet_hours(s: dict) -> bool:
    """Determina si el momento actual está en horario silencioso."""
    # Si el horario silencioso está desactivado, nunca es hora silenciosa
    if not s.get("quiet_enabled", True):
        return False
    import datetime as _dt2
    now_h = _dt2.datetime.now().hour
    now_d = _dt2.datetime.now().isoweekday()   # 1=Lun .. 7=Dom
    # Verificar día activo (si hoy no está en active_days → quiet)
    if now_d not in (s.get("active_days") or [1, 2, 3, 4, 5]):
        return True
    qs, qe = s.get("quiet_start", 22), s.get("quiet_end", 8)
    if qs > qe:   # cruza medianoche  (ej: 22-8)
        return now_h >= qs or now_h < qe
    return qs <= now_h < qe


@app.get("/api/whatsapp-personal/status")
async def wp_status(user: str = Depends(get_current_user)):
    """Estado del servicio + settings del usuario."""
    try:
        r = _http_client.get(f"{WHATSAPP_PERSONAL_URL}/status", timeout=5)
        bridge_data = r.json()
    except Exception:
        bridge_data = {"status": "unreachable", "phone": None, "hasQR": False}
    settings = _get_wp_settings(user)
    return {**bridge_data, "settings": settings}


@app.get("/api/whatsapp-personal/qr")
async def wp_qr(user: str = Depends(get_current_user)):
    """Devuelve el QR actual del servicio personal."""
    try:
        r = _http_client.get(f"{WHATSAPP_PERSONAL_URL}/qr-json", timeout=5)
        return r.json()
    except Exception as e:
        return {"status": "unreachable", "qr": None, "phone": None, "error": str(e)}


@app.get("/api/whatsapp-personal/settings")
async def wp_get_settings(user: str = Depends(get_current_user)):
    """Lee settings del usuario."""
    return _get_wp_settings(user)


@app.patch("/api/whatsapp-personal/settings")
async def wp_patch_settings(body: WhatsappPersonalSettings, user: str = Depends(get_current_user)):
    """Actualiza settings del usuario."""
    pool = get_postgres_pool()
    if not pool:
        raise HTTPException(503, "DB no disponible")
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            # Upsert
            cur.execute(
                "INSERT INTO whatsapp_personal_settings (user_id) VALUES (%s) ON CONFLICT (user_id) DO NOTHING",
                (user,)
            )
            updates, vals = [], []
            if body.auto_reply        is not None: updates.append("auto_reply=%s");        vals.append(body.auto_reply)
            if body.quiet_enabled     is not None: updates.append("quiet_enabled=%s");     vals.append(body.quiet_enabled)
            if body.quiet_start       is not None: updates.append("quiet_start=%s");       vals.append(body.quiet_start)
            if body.quiet_end         is not None: updates.append("quiet_end=%s");         vals.append(body.quiet_end)
            if body.active_days       is not None: updates.append("active_days=%s");       vals.append(body.active_days)
            if body.reply_scope       is not None: updates.append("reply_scope=%s");       vals.append(body.reply_scope)
            if body.allowed_contacts  is not None: updates.append("allowed_contacts=%s");  vals.append(json.dumps(body.allowed_contacts))
            if body.ai_assist         is not None: updates.append("ai_assist=%s");         vals.append(body.ai_assist)
            if body.offline_msg       is not None: updates.append("offline_msg=%s");       vals.append(body.offline_msg)
            if updates:
                updates.append("updated_at=NOW()")
                vals.append(user)
                cur.execute(f"UPDATE whatsapp_personal_settings SET {', '.join(updates)} WHERE user_id=%s", vals)
            conn.commit()
    finally:
        pool.putconn(conn)
    return _get_wp_settings(user)


@app.post("/api/whatsapp-personal/disconnect")
async def wp_disconnect(user: str = Depends(get_current_user)):
    """Desconecta la sesión WhatsApp personal."""
    try:
        r = _http_client.post(f"{WHATSAPP_PERSONAL_URL}/logout", timeout=10)
        return r.json()
    except Exception as e:
        raise HTTPException(503, f"No se pudo desconectar: {e}")


@app.post("/api/whatsapp-personal/connected")
async def wp_connected(request: Request):
    """Callback interno: el servicio personal avisa que se conectó."""
    secret = request.headers.get("X-Internal-Secret")
    if secret != INTERNAL_API_SECRET:
        raise HTTPException(403)
    data = await request.json()
    logging.info(f"[WA-PERSONAL] Conectado como: {data.get('phone')}")
    return {"ok": True}


@app.get("/api/whatsapp-personal/check-settings")
async def wp_check_settings(request: Request):
    """Endpoint interno: el servicio personal consulta settings antes de responder."""
    secret = request.headers.get("X-Internal-Secret")
    if secret != INTERNAL_API_SECRET:
        raise HTTPException(403)
    owner = os.environ.get("ADMIN_EMAIL", "ricardogs26@gmail.com")
    s = _get_wp_settings(owner)
    return {
        "auto_reply":        s["auto_reply"],
        "in_quiet_hours":    _is_in_quiet_hours(s),
        "reply_scope":       s["reply_scope"],
        "allowed_contacts":  s["allowed_contacts"],
        "ai_assist":         s["ai_assist"],
        "offline_msg":       s["offline_msg"],
    }


@app.get("/api/health")
async def health_check():
    return {"status": "ok"}